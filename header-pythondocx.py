from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.oxml.shared import qn as qn_shared

def set_font(run, size=13, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def set_line_spacing(paragraph, line=240):
    """Set line spacing - 240 = single line spacing in twips"""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

def add_vml_line(paragraph, start_x_pt=57, end_x_pt=228, y_pt=16, stroke_weight='0.5pt'):
    """Thêm đường gạch dưới dạng VML line (dùng trong văn bản hành chính)
    
    Parameters:
    - start_x_pt: tọa độ X bắt đầu (point)
    - end_x_pt: tọa độ X kết thúc (point) - càng nhỏ thì đường càng ngắn
    - y_pt: tọa độ Y (point) - càng nhỏ thì đường càng gần dòng trên
    - stroke_weight: độ dày đường (pt) - 0.5pt là mỏng, 1.5pt là dày
    """
    p = paragraph._p
    
    # Namespace VML
    vml_ns = 'urn:schemas-microsoft-com:vml'
    
    # Tạo run chứa VML
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    run.append(rPr)
    
    # Tạo pict element
    pict = OxmlElement('w:pict')
    
    # Tạo VML line với namespace đầy đủ
    line_xml = f'<v:line xmlns:v="{vml_ns}" style="position:absolute;z-index:251656704;visibility:visible;mso-wrap-style:square;mso-width-percent:0;mso-height-percent:0;mso-wrap-distance-left:9pt;mso-wrap-distance-top:0;mso-wrap-distance-right:9pt;mso-wrap-distance-bottom:0;mso-position-horizontal:absolute;mso-position-horizontal-relative:text;mso-position-vertical:absolute;mso-position-vertical-relative:text;mso-width-percent:0;mso-height-percent:0;mso-width-relative:page;mso-height-relative:page" from="{start_x_pt}pt,{y_pt}pt" to="{end_x_pt}pt,{y_pt}pt" strokecolor="#000000" strokeweight="{stroke_weight}"/>'
    
    line_element = parse_xml(line_xml)
    pict.append(line_element)
    run.append(pict)
    p.append(run)

# ====================== TẠO DOCUMENT ======================
doc = Document()

# Đặt khổ giấy A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(2)

# ====================== BẢNG 3 HÀNG X 2 CỘT ======================
table = doc.add_table(rows=3, cols=2)
table.autofit = False
table.columns[0].width = Cm(7.0)
table.columns[1].width = Cm(12.0)

# ==================== HÀNG 1 - CỘT TRÁI ====================
left_cell = table.cell(0, 0)

# UBND TỈNH ĐỒNG THÁP
p_ubnd = left_cell.add_paragraph()
p_ubnd.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing(p_ubnd, line=240)
run_ubnd = p_ubnd.add_run("UBND TỈNH ĐỒNG THÁP")
set_font(run_ubnd, size=13, bold=False)

# SỞ GIÁO DỤC VÀ ĐÀO TẠO
p_sogd = left_cell.add_paragraph()
p_sogd.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing(p_sogd, line=240)
run_sogd = p_sogd.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO")
set_font(run_sogd, size=13, bold=True)

# Đường gạch dưới SỞ GIÁO DỤC VÀ ĐÀO TẠO
p_line_sogd = left_cell.add_paragraph()
add_vml_line(p_line_sogd, start_x_pt=59, end_x_pt=160, y_pt=2, stroke_weight='0.75pt')

# ==================== HÀNG 1 - CỘT PHẢI ====================
right_cell = table.cell(0, 1)

# CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
p_qh = right_cell.add_paragraph()
p_qh.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing(p_qh, line=240)
run_qh = p_qh.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
set_font(run_qh, size=13, bold=True)

# Độc lập - Tự do - Hạnh phúc
p_dl = right_cell.add_paragraph()
p_dl.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing(p_dl, line=240)
run_dl = p_dl.add_run("Độc lập - Tự do - Hạnh phúc")
set_font(run_dl, size=14, bold=False)

# Đường gạch dưới Độc lập - Tự do - Hạnh phúc
p_line_dl = right_cell.add_paragraph()
add_vml_line(p_line_dl, start_x_pt=28, end_x_pt=190, y_pt=2, stroke_weight='0.75pt')

# ==================== HÀNG 2 - CỘT TRÁI ====================
left_cell_2 = table.cell(1, 0)
p_so = left_cell_2.add_paragraph()
p_so.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing(p_so, line=240)
run_so = p_so.add_run("Số:         /KH-SGDĐT")
set_font(run_so, size=13, bold=False)

# ==================== HÀNG 2 - CỘT PHẢI ====================
right_cell_2 = table.cell(1, 1)
p_date = right_cell_2.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing(p_date, line=240)
run_date = p_date.add_run("Đồng Tháp, ngày       tháng 4 năm 2026")
set_font(run_date, size=14, italic=True)

# ==================== HÀNG 3 - RỖNG ====================
# Hàng thứ 3 để trống (có thể dùng cho khoảng cách)

# ====================== LƯU FILE ======================
doc.save("VanBanHanhChinh_TEMPLATE_v7.docx")
print("Đã tạo file header theo mẫu TEMPLATE! Mở bằng Word để kiểm tra.")