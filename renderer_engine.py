#!/usr/bin/env python3
"""
Shared utilities dùng chung cho tất cả renderers.

Mỗi loại văn bản có renderer riêng:
    render_ke_hoach.py   → KẾ HOẠCH   (dùng TEMPLATE.docx)
    render_cong_van.py   → CÔNG VĂN   (dùng TEMPLATE_CV.docx)
    render_bao_cao.py    → BÁO CÁO    (chưa triển khai)
    render_quyet_dinh.py → QUYẾT ĐỊNH (chưa triển khai)
    render_to_trinh.py   → TỜ TRÌNH   (chưa triển khai)

Các renderer trên đều import từ file này.

Block Schema Version: 1.0
"""

import re
import json
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from pathlib import Path

# Block Schema Version - freeze for backward compatibility
BLOCK_SCHEMA_VERSION = "1.0"

# Block Registry - mapping block types to renderers
# This pattern avoids if/else spaghetti and allows easy extension
BLOCK_RENDERERS = {
    "heading1": None,  # Will be set after function definitions
    "heading2": None,
    "heading3": None,
    "paragraph": None,
    "bullet": None,
}

def register_block_renderer(block_type: str):
    """
    Decorator to register a block renderer
    
    Args:
        block_type: Type of block (heading1, heading2, etc.)
    """
    def decorator(func):
        BLOCK_RENDERERS[block_type] = func
        return func
    return decorator

def load_style_config(config_path: Optional[Path] = None) -> Dict:
    """
    Load style config từ file JSON
    
    Args:
        config_path: Đường dẫn file config (nếu None, dùng config mặc định)
    
    Returns:
        Dict với style config
    """
    if config_path is None:
        config_path = Path(__file__).parent / "style_config.json"
    
    if config_path.exists():
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    else:
        # Config mặc định
        return {
            "styles": {
                "heading1": {"bold": True, "font_size": 14, "indent_cm": 0},
                "heading2": {"bold": True, "font_size": 13, "indent_cm": 1.0},
                "heading3": {"bold": False, "font_size": 13, "indent_cm": 1.5},
                "paragraph": {"bold": False, "font_size": 13, "indent_cm": 1.25},
                "bullet": {"bold": False, "font_size": 13, "indent_cm": 1.25}
            },
            "font": {"name": "Times New Roman", "vietnamese_support": True},
            "paragraph": {"alignment": "justify"}
        }

def validate_blocks(blocks: List[Dict]) -> List[str]:
    """
    Validate structured blocks against schema
    
    Args:
        blocks: List of blocks to validate
    
    Returns:
        List of validation errors (empty if valid)
    
    Checks:
        - Missing required fields
        - Invalid block types
        - Invalid metadata
    """
    errors = []
    valid_types = {"heading1", "heading2", "heading3", "paragraph", "bullet"}
    
    for i, block in enumerate(blocks):
        # Check required fields
        if "type" not in block:
            errors.append(f"Block {i}: Missing 'type' field")
            continue
        
        if "text" not in block:
            errors.append(f"Block {i}: Missing 'text' field")
            continue
        
        # Check block type
        if block["type"] not in valid_types:
            errors.append(f"Block {i}: Invalid type '{block['type']}'. Valid types: {valid_types}")
        
        # Check metadata
        if "meta" not in block:
            errors.append(f"Block {i}: Missing 'meta' field")
        elif not isinstance(block["meta"], dict):
            errors.append(f"Block {i}: 'meta' must be a dict")
    
    return errors

def parse_content_to_blocks(content: str) -> List[Dict]:
    """
    Parse nội dung text thành structured blocks với metadata
    
    Args:
        content: Nội dung text từ AI
    
    Returns:
        List of blocks với schema chuẩn:
        {
            "type": "heading1|heading2|heading3|paragraph|bullet",
            "text": "nội dung",
            "meta": {
                "numbering": true/false,
                "indent": true/false,
                ... (future extensibility)
            }
        }
    """
    lines = content.split('\n')
    blocks = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Heading level 1: I, II, III, IV, V...
        if re.match(r'^[IVXLCDM]+\.', line):
            blocks.append({
                'type': 'heading1',
                'text': line,
                'meta': {
                    'numbering': True,
                    'indent': False
                }
            })
        # Heading level 2: 1, 2, 3...
        # Chỉ coi là heading khi là dòng tiêu đề ngắn.
        # Các dòng đánh số dạng câu đầy đủ sẽ giữ là paragraph để không bị in đậm sai.
        elif re.match(r'^\d+\.', line):
            after_number = re.sub(r'^\d+\.\s*', '', line).strip()
            is_short_heading = len(after_number) <= 90 and not after_number.endswith('.')
            block_type = 'heading2' if is_short_heading else 'paragraph'
            blocks.append({
                'type': block_type,
                'text': line,
                'meta': {
                    'numbering': True,
                    'indent': True
                }
            })
        # Heading level 3: a, b, c...
        elif re.match(r'^[a-z]+\.', line):
            blocks.append({
                'type': 'heading3',
                'text': line,
                'meta': {
                    'numbering': True,
                    'indent': True
                }
            })
        # Bullet
        elif line.startswith('• '):
            blocks.append({
                'type': 'bullet',
                'text': line,
                'meta': {
                    'numbering': True,
                    'indent': True
                }
            })
        # Normal paragraph
        else:
            blocks.append({
                'type': 'paragraph',
                'text': line,
                'meta': {
                    'numbering': False,
                    'indent': True
                }
            })
    
    return blocks

def set_font(run, bold=False, size=13, style_config=None):
    """
    Set font cho run dựa trên style config
    
    Args:
        run: Run object
        bold: Bold flag
        size: Font size in points
        style_config: Style config dict (nếu None, dùng mặc định)
    """
    if style_config is None:
        style_config = load_style_config()
    
    font_config = style_config.get('font', {})
    font_name = font_config.get('name', 'Times New Roman')
    
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold

    # Set font tiếng Việt nếu config yêu cầu
    if font_config.get('vietnamese_support', True):
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)


def add_vml_line(paragraph, start_x_pt=0, end_x_pt=85, y_pt=2, stroke_weight='0.75pt'):
    """Thêm đường line VML như mẫu hành chính (độ dài ~3cm)."""
    p = paragraph._p
    vml_ns = 'urn:schemas-microsoft-com:vml'

    run = OxmlElement('w:r')
    run.append(OxmlElement('w:rPr'))

    pict = OxmlElement('w:pict')
    line_xml = (
        f'<v:line xmlns:v="{vml_ns}" '
        'style="position:absolute;z-index:251656704;visibility:visible;'
        'mso-wrap-style:square;mso-width-percent:0;mso-height-percent:0;'
        'mso-wrap-distance-left:9pt;mso-wrap-distance-top:0;'
        'mso-wrap-distance-right:9pt;mso-wrap-distance-bottom:0;'
        'mso-position-horizontal:center;mso-position-horizontal-relative:text;'
        'mso-position-vertical:absolute;mso-position-vertical-relative:text;'
        'mso-width-relative:page;mso-height-relative:page" '
        f'from="{start_x_pt}pt,{y_pt}pt" to="{end_x_pt}pt,{y_pt}pt" '
        f'strokecolor="#000000" strokeweight="{stroke_weight}"/>'
    )
    pict.append(parse_xml(line_xml))
    run.append(pict)
    p.append(run)

def set_paragraph_format(para, block_type='paragraph', style_config=None):
    """
    Set format cho paragraph dựa trên style config
    
    Args:
        para: Paragraph object
        block_type: Type của block (heading1, heading2, heading3, paragraph, bullet)
        style_config: Style config dict (nếu None, dùng mặc định)
    """
    if style_config is None:
        style_config = load_style_config()
    
    # Get style cho block type
    styles = style_config.get('styles', {})
    block_style = styles.get(block_type, styles.get('paragraph', {}))
    
    # Set alignment
    para_config = style_config.get('paragraph', {})
    alignment = para_config.get('alignment', 'justify')
    if alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'left':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Set line spacing
    # - value <= 3: line multiple/single spacing
    # - value > 3: exact point spacing
    line_spacing = block_style.get('line_spacing', 1.0)
    if isinstance(line_spacing, (int, float)) and line_spacing <= 3:
        para.paragraph_format.line_spacing = line_spacing
    else:
        para.paragraph_format.line_spacing = Pt(line_spacing)
    
    # Set spacing after
    spacing_after = block_style.get('spacing_after', 3)
    para.paragraph_format.space_after = Pt(spacing_after)
    
    # Set spacing before
    spacing_before = block_style.get('spacing_before', 0)
    if spacing_before > 0:
        para.paragraph_format.space_before = Pt(spacing_before)
    
    # Set thụt đầu dòng (sử dụng Inches để chính xác hơn)
    indent_cm = block_style.get('indent_cm', 0)
    para.paragraph_format.first_line_indent = Inches(indent_cm / 2.54)  # Convert cm to inches



def _normalize_default_header_blank_lines(doc: Document) -> None:
    """Xóa toàn bộ paragraph rỗng ở cuối header mặc định (áp dụng từ trang 2)."""
    for section in doc.sections:
        header = section.header
        paragraphs = list(header.paragraphs)
        if not paragraphs:
            continue

        last_non_empty_idx = -1
        for i, para in enumerate(paragraphs):
            if (para.text or '').strip():
                last_non_empty_idx = i

        # Nếu header toàn dòng trống thì không đụng để tránh phá layout không mong muốn.
        if last_non_empty_idx == -1:
            continue

        for para in paragraphs[last_non_empty_idx + 1:]:
            if (para.text or '').strip() == '':
                p = para._element
                p.getparent().remove(p)

def clean_content(content: str) -> str:
    """
    Lọc bỏ tiêu đề và footer từ nội dung AI
    
    Args:
        content: Nội dung thô từ AI
    
    Returns:
        Nội dung đã làm sạch
    """
    lines = content.split('\n')
    
    # Bỏ dòng đầu tiên nếu là loại văn bản
    if lines and lines[0].strip().upper() in ["KẾ HOẠCH", "QUYẾT ĐỊNH", "BÁO CÁO", "TỜ TRÌNH", "BIÊN BẢN", "CÔNG VĂN"]:
        lines = lines[1:]
    
    # Bỏ dòng trích yếu ngắn (nếu có), nhưng giữ lại các dòng mở đầu kiểu "Căn cứ..."
    if lines and lines[0].strip():
        first_line = lines[0].strip()
        first_upper = first_line.upper()
        looks_like_legal_basis = first_upper.startswith("CĂN CỨ")
        if not looks_like_legal_basis and len(first_line) <= 220:
            lines = lines[1:]
    
    # Bỏ dòng trống tiếp theo
    if lines and not lines[0].strip():
        lines = lines[1:]
    
    # Lọc nội dung chính, bỏ footer
    result = []
    for line in lines:
        line = line.strip()
        
        # Tìm dấu hiệu footer
        if "Nơi nhận:" in line or re.match(r"^[^,]+,\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+", line, re.IGNORECASE):
            break
        
        # Loại bỏ "./." ở cuối
        line = line.replace("./.", "").strip()
        
        if line:
            result.append(line)
        else:
            result.append('')
    
    return '\n'.join(result)


def _format_noi_nhan_table(doc: Document, noi_nhan_text: str, style_config: Optional[Dict] = None) -> None:
    """Định dạng bảng Nơi nhận:
    - "Nơi nhận:" size 12
    - Các dòng bên dưới size 11
    - Căn trái để tránh giãn chữ do justify
    """
    if style_config is None:
        style_config = load_style_config()

    target_table = None
    for table in doc.tables:
        found = False
        for row in table.rows:
            for cell in row.cells:
                if 'Nơi nhận' in (cell.text or ''):
                    found = True
                    break
            if found:
                break
        if found:
            target_table = table
            break

    if target_table is None:
        return

    left_cell = target_table.cell(0, 0)
    left_cell.text = ""

    p_header = left_cell.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(0)
    p_header.paragraph_format.line_spacing = 1.0
    p_header.paragraph_format.first_line_indent = Inches(0)
    run_header = p_header.add_run("Nơi nhận:")
    set_font(run_header, bold=False, size=12, style_config=style_config)
    run_header.font.bold = True
    run_header.font.italic = True

    lines = [ln.strip() for ln in (noi_nhan_text or '').splitlines() if ln.strip()]
    for line in lines:
        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(line)
        set_font(run, bold=False, size=11, style_config=style_config)

