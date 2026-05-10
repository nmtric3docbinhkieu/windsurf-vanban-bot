#!/usr/bin/env python3
"""
Script final: docxtpl inject text thường + python-docx format lại đề mục
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate
from pathlib import Path
import re

def clean_content(content):
    """Lọc bỏ tiêu đề và footer"""
    lines = content.split('\n')
    
    # Bỏ dòng đầu tiên nếu là loại văn bản
    if lines and lines[0].strip().upper() in ["KẾ HOẠCH", "QUYẾT ĐỊNH", "BÁO CÁO", "TỜ TRÌNH", "BIÊN BẢN", "CÔNG VĂN"]:
        lines = lines[1:]
    
    # Bỏ dòng thứ hai nếu là trích yếu
    if lines and lines[0].strip():
        lines = lines[1:]
    
    # Bỏ dòng trống tiếp theo
    if lines and not lines[0].strip():
        lines = lines[1:]
    
    # Lọc nội dung chính, bỏ footer
    result = []
    for line in lines:
        line = line.strip()
        
        # Tìm dấu hiệu footer
        if "Nơi nhận:" in line or re.match(r"^[^,]+,\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+", line, re.IGNORECASE):
            break
        
        # Loại bỏ "./." ở cuối
        line = line.replace("./.", "").strip()
        
        if line:
            result.append(line)
        else:
            result.append('')
    
    return '\n'.join(result)

def format_paragraph(para, is_heading=False, indent_level=0):
    """Format paragraph"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = Pt(18)
    para.paragraph_format.space_after = Pt(3)
    
    if indent_level > 0:
        if indent_level == 1:
            para.paragraph_format.first_line_indent = Cm(1.0)
        elif indent_level == 2:
            para.paragraph_format.first_line_indent = Cm(1.5)
        elif indent_level == 3:
            para.paragraph_format.first_line_indent = Cm(1.25)
    
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.bold = is_heading
        
        # Set font tiếng Việt
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rPr.append(rFonts)

def format_document(doc):
    """Format lại document: split nội dung, in đậm đề mục, thụt đầu dòng"""
    new_paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            new_paragraphs.append(para)
            continue
        
        # Nếu nội dung có \n (nhiều dòng), split thành các paragraph riêng
        if '\n' in text and not ('{{' in text or '}}' in text):
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # Tạo paragraph mới cho mỗi dòng
                    new_para = doc.add_paragraph(line)
                    new_paragraphs.append(new_para)
                else:
                    new_paragraphs.append(doc.add_paragraph())
        else:
            new_paragraphs.append(para)
    
    # Xóa các paragraph cũ, giữ lại các paragraph mới
    # Cách đơn giản: xóa tất cả, thêm lại
    # Nhưng sẽ mất header/footer trong table
    
    # Cách khác: chỉ format lại các paragraph
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Đề mục cấp I, II, III - in đậm, không thụt đầu dòng
        if re.match(r'^[IVX]+\.', text):
            format_paragraph(para, is_heading=True, indent_level=0)
        # Đề mục cấp 1, 2, 3 - in đậm, thụt đầu dòng 1cm
        elif re.match(r'^\d+\.', text):
            format_paragraph(para, is_heading=True, indent_level=1)
        # Đề mục cấp a, b, c - bình thường, thụt đầu dòng 1.5cm
        elif re.match(r'^[a-z]+\.', text):
            format_paragraph(para, is_heading=False, indent_level=2)
        # Bullet
        elif text.startswith('• '):
            format_paragraph(para, is_heading=False, indent_level=3)
        # Dòng thường - thụt đầu dòng 1.25cm
        else:
            # Không format nếu là placeholder hoặc header/footer
            if '{{' not in text and '}}' not in text:
                format_paragraph(para, is_heading=False, indent_level=3)

def split_multiline_content(doc):
    """Split các paragraph có \n thành nhiều paragraph riêng và format luôn"""
    paragraphs_to_process = []
    
    for para in doc.paragraphs:
        if '\n' in para.text and not ('{{' in para.text or '}}' in para.text):
            paragraphs_to_process.append(para)
    
    for para in paragraphs_to_process:
        text = para.text
        lines = text.split('\n')
        
        # Lưu vị trí của paragraph hiện tại
        parent = para._element.getparent()
        index = parent.index(para._element)
        
        # Xóa paragraph cũ
        parent.remove(para._element)
        
        # Thêm các paragraph mới và format luôn
        for i, line in enumerate(lines):
            line = line.strip()
            if line:
                new_para = doc.add_paragraph(line)
                # Format ngay khi tạo
                if re.match(r'^[IVX]+\.', line):
                    format_paragraph(new_para, is_heading=True, indent_level=0)
                elif re.match(r'^\d+\.', line):
                    format_paragraph(new_para, is_heading=True, indent_level=1)
                elif re.match(r'^[a-z]+\.', line):
                    format_paragraph(new_para, is_heading=False, indent_level=2)
                else:
                    format_paragraph(new_para, is_heading=False, indent_level=3)
                
                new_elem = new_para._element
                parent.insert(index + i, new_elem)
            else:
                new_para = doc.add_paragraph()
                new_elem = new_para._element
                parent.insert(index + i, new_elem)

def create_van_ban_final(template_path, content, output_path, metadata):
    """
    Tạo văn bản: docxtpl inject text thường + python-docx format lại
    
    Args:
        template_path: File template có placeholder
        content: Nội dung chính
        output_path: File output
        metadata: Dict với so_ky_hieu, ngay_thang, trich_yeu, noi_nhan, nguoi_ky, chuc_vu_ky
    """
    # Bước 1: Inject placeholder bằng docxtpl với text thường
    doc_tpl = DocxTemplate(template_path)
    
    content = clean_content(content)
    
    data = {
        'loai_van_ban': 'KẾ HOẠCH',
        'so_ky_hieu': metadata.get('so_ky_hieu', ''),
        'ngay_thang': metadata.get('ngay_thang', ''),
        'trich_yeu': metadata.get('trich_yeu', ''),
        'noi_dung': content,  # Inject text thường
        'noi_nhan': metadata.get('noi_nhan', ''),
        'nguoi_ky': metadata.get('nguoi_ky', ''),
        'chuc_vu_ky': metadata.get('chuc_vu_ky', '')
    }
    
    doc_tpl.render(data)
    doc_tpl.save(output_path)
    
    # Bước 2: Mở lại bằng python-docx để split và format lại đề mục
    doc = Document(output_path)
    split_multiline_content(doc)
    format_document(doc)  # Format lại tất cả paragraph
    doc.save(output_path)
    
    print(f"✅ Đã tạo file: {output_path}")

def test():
    """Test với dữ liệu mẫu"""
    template_path = Path(__file__).parent / "TEMPLATE.docx"
    content_file = Path(__file__).parent / "noi_dung_ke_hoach.txt"
    output_path = Path(__file__).parent.parent / "van-ban-di" / "test_final_v4.docx"
    
    content = content_file.read_text(encoding='utf-8')
    
    metadata = {
        'so_ky_hieu': '123/KH-THPTĐBK',
        'ngay_thang': 'Đốc Bình Kiều, ngày 10 tháng 5 năm 2026',
        'trich_yeu': 'Triển khai thực hiện Chỉ thị số 03/CT-TTg về ngăn chặn bạo lực học đường',
        'noi_nhan': 'Sở GDĐT Đồng Tháp (báo cáo); Lưu: VT',
        'nguoi_ky': 'Nguyễn Minh Trí',
        'chuc_vu_ky': 'KT. HIỆU TRƯỞNG\nPHÓ HIỆU TRƯỞNG'
    }
    
    create_van_ban_final(template_path, content, output_path, metadata)

if __name__ == "__main__":
    test()
