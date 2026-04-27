#!/usr/bin/env python3
"""
Script chuẩn hóa văn bản hành chính - BẢN HOÀN CHỈNH
Đường gạch ngang sát chữ, đúng thể thức
"""

import argparse
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# ============================================================
# CẤU HÌNH THỂ THỨC (CÓ THỂ SỬA THEO TRƯỜNG BẠN)
# ============================================================
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(14)
FONT_SIZE_TITLE = Pt(14)
FONT_SIZE_SUBTITLE = Pt(14)
FONT_SIZE_HEADER = Pt(13)
FONT_SIZE_SIGN = Pt(14)
FONT_SIZE_NHAN = Pt(12)
FONT_SIZE_NHAN_CONTENT = Pt(11)

LINE_SPACING_PT = Pt(18)
FIRST_LINE_INDENT = Cm(1.27)
PARA_SPACE_AFTER = Pt(3)

LEFT_MARGIN = Cm(3.0)
RIGHT_MARGIN = Cm(2.0)
TOP_MARGIN = Cm(2.0)
BOTTOM_MARGIN = Cm(2.0)

CO_QUAN_CHU_QUAN = "SỞ GDĐT ĐỒNG THÁP"
CO_QUAN_BAN_HANH = "TRƯỜNG THPT ĐỐC BINH KIỀU"
TEN_NGUOI_KY = "Nguyễn Minh Trí"
CHUC_VU_KY = "PHÓ HIỆU TRƯỞNG"
KT_HIEU_TRUONG = "KT. HIỆU TRƯỞNG"
DIA_DANH = "Đốc Binh Kiều"

# ============================================================
# HÀM TIỆN ÍCH ĐỊNH DẠNG
# ============================================================
def set_font(run, name=FONT_NAME, size=FONT_SIZE_BODY, bold=False, italic=False):
    run.font.name = name
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = size
    
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rPr.append(rFonts)

def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
                  size=FONT_SIZE_BODY, space_after=PARA_SPACE_AFTER, first_line_indent=FIRST_LINE_INDENT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = LINE_SPACING_PT
    pf.first_line_indent = first_line_indent
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p

def add_centered_paragraph(doc, text, bold=False, size=FONT_SIZE_BODY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def make_invisible_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_vml_line(paragraph, start_x_pt=57, end_x_pt=228, y_pt=16, stroke_weight='0.5pt'):
    """Thêm đường gạch dưới dạng VML line (dùng trong văn bản hành chính)
    
    Parameters:
    - start_x_pt: tọa độ X bắt đầu (point)
    - end_x_pt: tọa độ X kết thúc (point) - càng nhỏ thì đường càng ngắn
    - y_pt: tọa độ Y (point) - càng nhỏ thì đường càng gần dòng trên
    - stroke_weight: độ dày đường (pt) - 0.5pt là mỏng, 1.5pt là dày
    """
    p = paragraph._p
    
    # Namespace VML
    vml_ns = 'urn:schemas-microsoft-com:vml'
    
    # Tạo run chứa VML
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    run.append(rPr)
    
    # Tạo pict element
    pict = OxmlElement('w:pict')
    
    # Tạo VML line với namespace đầy đủ
    line_xml = f'<v:line xmlns:v="{vml_ns}" style="position:absolute;z-index:251656704;visibility:visible;mso-wrap-style:square;mso-width-percent:0;mso-height-percent:0;mso-wrap-distance-left:9pt;mso-wrap-distance-top:0;mso-wrap-distance-right:9pt;mso-wrap-distance-bottom:0;mso-position-horizontal:absolute;mso-position-horizontal-relative:text;mso-position-vertical:absolute;mso-position-vertical-relative:text;mso-width-percent:0;mso-height-percent:0;mso-width-relative:page;mso-height-relative:page" from="{start_x_pt}pt,{y_pt}pt" to="{end_x_pt}pt,{y_pt}pt" strokecolor="#000000" strokeweight="{stroke_weight}"/>'
    
    line_element = parse_xml(line_xml)
    pict.append(line_element)
    run.append(pict)
    p.append(run)

def add_line_using_table(cell, col_width_dxa):
    """Vẽ đường gạch bằng table lồng - nằm trong cell chiều cao rất nhỏ"""
    # Tạo bảng con 1x1
    tbl = cell.add_table(rows=1, cols=1)
    tbl.autofit = False
    
    # Chiều rộng bảng = 50% cell
    line_width = int(col_width_dxa * 0.5)
    tblPr = tbl._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(line_width))
    tblPr.append(tblW)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cell bên trong
    inner_cell = tbl.cell(0, 0)
    # Set chiều cao cell rất nhỏ
    tcPr = inner_cell._tc.get_or_add_tcPr()
    tcHeight = OxmlElement('w:tcH')
    tcHeight.set(qn('w:val'), '1')
    tcPr.append(tcHeight)
    
    # Paragraph với space = 0
    p = inner_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    
    # Đường kẻ bottom
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

def _set_cell_width_dxa(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(dxa))
    tcPr.insert(0, tcW)

def _set_table_fixed_width(tbl, total_dxa, indent_dxa=0):
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tblEl.insert(0, tblPr)
    
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(total_dxa))
    tblPr.insert(0, tblW)
    
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')
    
    if indent_dxa != 0:
        for old in tblPr.findall(qn('w:tblInd')):
            tblPr.remove(old)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:type'), 'dxa')
        tblInd.set(qn('w:w'), str(indent_dxa))
        tblPr.append(tblInd)
    tbl.autofit = False

def create_header(doc, so_ky_hieu, ngay_thang_nam=None):
    if ngay_thang_nam is None:
        ngay_thang_nam = f"{DIA_DANH}, ngày    tháng    năm 2026"
    
    COL0_DXA = 4500  # 7.8cm
    COL1_DXA = 5670  # 10cm
    TOTAL_DXA = 10170   
    INDENT_DXA = -500
    
    table = doc.add_table(rows=3, cols=2)
    for row in table.rows:
        for cell in row.cells:
            make_invisible_border(cell)
    
    _set_table_fixed_width(table, TOTAL_DXA, INDENT_DXA)
    _set_cell_width_dxa(table.cell(0, 0), COL0_DXA)
    _set_cell_width_dxa(table.cell(0, 1), COL1_DXA)
    
    # === Cột trái: Row 0 (chữ) ===
    cell_left = table.cell(0, 0)
    p1 = cell_left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing = 1.0
    run1 = p1.add_run(CO_QUAN_CHU_QUAN)
    set_font(run1, size=FONT_SIZE_HEADER, bold=False)
    
    p2 = cell_left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    run2 = p2.add_run(CO_QUAN_BAN_HANH)
    set_font(run2, size=FONT_SIZE_HEADER, bold=True)
    
    # === Cột trái: Row 0 (đường gạch VML) ===
    p_line_truongdbk = cell_left.add_paragraph()
    p_line_truongdbk.paragraph_format.space_before = Pt(0)
    p_line_truongdbk.paragraph_format.space_after = Pt(0)
    p_line_truongdbk.paragraph_format.line_spacing = Pt(1)
    add_vml_line(p_line_truongdbk, start_x_pt=59, end_x_pt=160, y_pt=2, stroke_weight='0.75pt')
    
    # === Cột phải: Row 0 (chữ) ===
    cell_right = table.cell(0, 1)
    pr1 = cell_right.paragraphs[0]
    pr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr1.paragraph_format.space_before = Pt(0)
    pr1.paragraph_format.space_after = Pt(0)
    pr1.paragraph_format.line_spacing = 1.0
    rr1 = pr1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(rr1, size=FONT_SIZE_HEADER, bold=True)
    
    pr2 = cell_right.add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr2.paragraph_format.space_before = Pt(0)
    pr2.paragraph_format.space_after = Pt(0)
    pr2.paragraph_format.line_spacing = 1.0
    rr2 = pr2.add_run("Độc lập – Tự do – Hạnh phúc")
    set_font(rr2, size=Pt(14), bold=True)
    
    # === Cột phải: Row 0 (đường gạch VML) ===
    p_line_dl = cell_right.add_paragraph()
    p_line_dl.paragraph_format.space_before = Pt(0)
    p_line_dl.paragraph_format.space_after = Pt(0)
    p_line_dl.paragraph_format.line_spacing = Pt(1)
    add_vml_line(p_line_dl, start_x_pt=50, end_x_pt=224, y_pt=2, stroke_weight='0.75pt')
    
    # === Row 1: Số ký hiệu + Ngày tháng ===
    cell_so = table.cell(1, 0)
    p_so = cell_so.paragraphs[0]
    p_so.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_so.paragraph_format.space_before = Pt(8)
    p_so.paragraph_format.space_after = Pt(0)
    run_so = p_so.add_run(f"Số: {so_ky_hieu}")
    set_font(run_so, size=FONT_SIZE_HEADER, bold=False)
    
    cell_ngay = table.cell(1, 1)
    p_ngay = cell_ngay.paragraphs[0]
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ngay.paragraph_format.space_before = Pt(8)
    p_ngay.paragraph_format.space_after = Pt(0)
    run_ngay = p_ngay.add_run(ngay_thang_nam)
    set_font(run_ngay, size=FONT_SIZE_HEADER, italic=True)

def create_title(doc, loai_van_ban, trich_yeu):
    loai_map = {
        "ke_hoach": "KẾ HOẠCH",
        "bao_cao": "BÁO CÁO",
        "to_trinh": "TỜ TRÌNH",
        "quyet_dinh": "QUYẾT ĐỊNH",
        "cong_van": "CÔNG VĂN"
    }
    loai_vb = loai_map.get(loai_van_ban, "VĂN BẢN")
    if loai_van_ban != "cong_van":
        doc.add_paragraph()
        add_centered_paragraph(doc, loai_vb, bold=True, size=FONT_SIZE_TITLE)
        add_centered_paragraph(doc, trich_yeu, bold=True, size=FONT_SIZE_SUBTITLE)
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"V/v {trich_yeu}")
        set_font(run, size=FONT_SIZE_BODY)
        doc.add_paragraph()

def add_markdown_content(doc, text):
    lines = text.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if (line.startswith("**") and line.endswith("**")) or \
           re.match(r"^[IVX]+\.", line) or re.match(r"^\d+\.", line):
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            add_paragraph(doc, clean, bold=True, space_after=Pt(3))
            continue
        if re.match(r"^[-*+]\s+", line):
            clean = re.sub(r"^[-*+]\s+", "", line)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(f"- {clean}")
            set_font(run, size=FONT_SIZE_BODY)
            continue
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        add_paragraph(doc, clean)

def create_receiver_and_signature(doc, noi_nhan_raw):
    table = doc.add_table(rows=1, cols=2)
    for cell in table.rows[0].cells:
        make_invisible_border(cell)
    
    cell_left = table.cell(0, 0)
    p_title = cell_left.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    p_title.paragraph_format.line_spacing = 1.0
    run_title = p_title.add_run("Nơi nhận:")
    set_font(run_title, size=FONT_SIZE_NHAN, bold=True, italic=True)
    
    receivers = [x.strip() for x in noi_nhan_raw.split(";") if x.strip()]
    for r in receivers[:6]:
        p = cell_left.add_paragraph()
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(f"- {r};")
        set_font(run, size=FONT_SIZE_NHAN_CONTENT)
    
    cell_right = table.cell(0, 1)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p_kt = cell_right.paragraphs[0]
    p_kt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kt.paragraph_format.space_before = Pt(0)
    p_kt.paragraph_format.space_after = Pt(0)
    p_kt.paragraph_format.line_spacing = 1.0
    run_kt = p_kt.add_run(KT_HIEU_TRUONG)
    set_font(run_kt, size=FONT_SIZE_SIGN, bold=True)
    
    p_cv = cell_right.add_paragraph()
    p_cv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cv.paragraph_format.space_before = Pt(0)
    p_cv.paragraph_format.space_after = Pt(0)
    p_cv.paragraph_format.line_spacing = 1.0
    run_cv = p_cv.add_run(CHUC_VU_KY)
    set_font(run_cv, size=FONT_SIZE_SIGN, bold=True)
    
    for _ in range(3):
        p_empty = cell_right.add_paragraph()
        p_empty.paragraph_format.space_before = Pt(0)
        p_empty.paragraph_format.space_after = Pt(0)
        p_empty.paragraph_format.line_spacing = 1.0
    
    p_ten = cell_right.add_paragraph()
    p_ten.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ten.paragraph_format.space_before = Pt(0)
    p_ten.paragraph_format.space_after = Pt(0)
    p_ten.paragraph_format.line_spacing = 1.0
    run_ten = p_ten.add_run(TEN_NGUOI_KY)
    set_font(run_ten, size=FONT_SIZE_SIGN, bold=True)

def read_text_from_file(file_path):
    path = Path(file_path)
    if path.suffix.lower() == '.txt':
        return path.read_text(encoding='utf-8')
    elif path.suffix.lower() == '.docx':
        doc = Document(path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paras)
    else:
        raise ValueError("Chỉ hỗ trợ file .txt hoặc .docx")

def tao_van_ban_chuan(input_file, output_file, loai, so_ky_hieu, trich_yeu, noi_nhan, ngay_thang=None):
    noi_dung_tho = read_text_from_file(input_file)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    
    create_header(doc, so_ky_hieu, ngay_thang)
    create_title(doc, loai, trich_yeu)
    add_markdown_content(doc, noi_dung_tho)
    create_receiver_and_signature(doc, noi_nhan)
    
    doc.save(output_file)
    print(f"✅ Đã tạo file: {output_file}")

def main():
    parser = argparse.ArgumentParser(description="Chuẩn hóa văn bản thô thành file docx đúng thể thức hành chính")
    parser.add_argument("--input", required=True, help="File đầu vào (txt hoặc docx)")
    parser.add_argument("--output", required=True, help="File đầu ra (.docx)")
    parser.add_argument("--loai", required=True, choices=["cong_van", "ke_hoach", "bao_cao", "to_trinh", "quyet_dinh"])
    parser.add_argument("--so", required=True, help="Số ký hiệu")
    parser.add_argument("--trich_yeu", required=True, help="Trích yếu")
    parser.add_argument("--noi_nhan", required=True, help="Danh sách nơi nhận, cách nhau ;")
    parser.add_argument("--ngay", help="Dòng ngày tháng")
    args = parser.parse_args()
    tao_van_ban_chuan(args.input, args.output, args.loai, args.so, args.trich_yeu, args.noi_nhan, args.ngay)

if __name__ == "__main__":
    main()