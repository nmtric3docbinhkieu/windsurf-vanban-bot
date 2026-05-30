#!/usr/bin/env python3
"""Kiểm tra nội dung file docx output"""

from docx import Document
from pathlib import Path

project_root = Path(__file__).resolve().parents[2]
output_file = project_root / "van-ban-di" / "test_renderer_v3.docx"
doc = Document(output_file)

print("=== TOÀN BỘ NỘI DUNG FILE OUTPUT ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text or i < 20:  # Hiện cả dòng trống
        print(f"{i:3d}: [{repr(text)}]")

print("\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i}: {len(table.rows)} rows, {len(table.columns)} cols ---")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  [{r_idx},{c_idx}]: {text[:50]}")
