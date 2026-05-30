#!/usr/bin/env python3
"""Kiểm tra placeholder trong TEMPLATE.docx"""

from docx import Document
from pathlib import Path

template_path = Path(__file__).resolve().parents[1] / "TEMPLATE.docx"
doc = Document(template_path)

print("=== NỘI DUNG TEMPLATE.DOCX ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text or i < 20:
        print(f"{i:3d}: [{repr(text)}]")

print("\n=== PLACEHOLDER TÌM ĐƯỢC ===")
placeholders = []
for para in doc.paragraphs:
    if '{{' in para.text and '}}' in para.text:
        placeholders.append(para.text)

if placeholders:
    for p in placeholders:
        print(f"  - {p}")
else:
    print("  ⚠️  Không tìm thấy placeholder nào trong TEMPLATE.docx")
    print("  Bạn cần mở TEMPLATE.docx trong Word và chèn placeholder:")
    print("  - {{so_ky_hieu}} - Số ký hiệu")
    print("  - {{ngay_thang}} - Ngày tháng")
    print("  - {{trich_yeu}} - Trích yếu")
    print("  - {{noi_dung}} - Nội dung chính (QUAN TRỌNG)")
    print("  - {{noi_nhan}} - Nơi nhận")
    print("  - {{nguoi_ky}} - Người ký")
    print("  - {{chuc_vu_ky}} - Chức vụ ký")
