#!/usr/bin/env python3
"""
Renderer dành riêng cho thể loại văn bản KẾ HOẠCH.

Chịu trách nhiệm duy nhất: render một KẾ HOẠCH từ structured blocks
vào file TEMPLATE.docx theo đúng thể thức hành chính.

Template sử dụng: TEMPLATE.docx

Kiến trúc tổng thể:
    soan_van_ban_di.py   → AI orchestration (đọc file, gọi AI, phân loại)
    render_ke_hoach.py   → Renderer KẾ HOẠCH  (file này)
    render_cong_van.py   → Renderer CÔNG VĂN
    render_bao_cao.py    → Renderer BÁO CÁO    (chưa triển khai)
    render_quyet_dinh.py → Renderer QUYẾT ĐỊNH (chưa triển khai)
    render_to_trinh.py   → Renderer TỜ TRÌNH   (chưa triển khai)
    renderer_engine.py   → Shared utilities (dùng chung cho tất cả renderers)

API chính:
    render_document(template_path, output_path, metadata, blocks, style_config_path=None)
"""

import re
from pathlib import Path
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

from renderer_engine import (
    add_vml_line,
    clean_content,
    load_style_config,
    parse_content_to_blocks,
    set_font,
    set_paragraph_format,
    validate_blocks,
    _format_noi_nhan_table,
    _normalize_default_header_blank_lines,
)

DEFAULT_TEMPLATE = Path(__file__).parent / "TEMPLATE.docx"


# ---------------------------------------------------------------------------
# Hàm nội bộ - chỉ dùng bởi render_document()
# ---------------------------------------------------------------------------

def _is_empty_paragraph_element(element) -> bool:
    """Kiểm tra phần tử body có phải paragraph rỗng hay không."""
    if not element.tag.endswith('p'):
        return False
    return ''.join(element.itertext()).strip() == ''


def _add_title_separator_line(doc: Document, style_config: Optional[Dict] = None) -> None:
    """Vẽ đường line VML ngắn (~3 cm) căn giữa sát dưới dòng tiêu đề phụ."""
    if style_config is None:
        style_config = load_style_config()

    non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty_paras) < 2:
        return

    subtitle_para = non_empty_paras[1]
    next_para = subtitle_para._p.getnext()
    if next_para is not None:
        next_xml = next_para.xml
        if '<w:pict' in next_xml and 'urn:schemas-microsoft-com:vml' in next_xml:
            return

    line_para = doc.add_paragraph('')
    subtitle_para._p.addnext(line_para._p)
    subtitle_para.paragraph_format.space_after = Pt(0)
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(1)
    line_para.paragraph_format.line_spacing = 1.0
    line_para.paragraph_format.first_line_indent = Inches(0)
    add_vml_line(line_para, start_x_pt=0, end_x_pt=85, y_pt=2, stroke_weight='0.75pt')


def _format_date_line(doc: Document, style_config: Optional[Dict] = None) -> None:
    """Chuẩn hóa dòng ngày tháng ở header phải về cỡ chữ 13, in nghiêng."""
    if style_config is None:
        style_config = load_style_config()

    # Vị trí chuẩn trong TEMPLATE.docx: bảng đầu tiên, hàng 2, cột phải.
    if doc.tables:
        try:
            date_cell = doc.tables[0].cell(1, 1)
            for para in date_cell.paragraphs:
                if not (para.text or '').strip():
                    continue
                for run in para.runs:
                    if run.text.strip():
                        set_font(run, bold=bool(run.bold), size=13, style_config=style_config)
                        run.font.italic = True
                return
        except Exception:
            pass

    # Fallback: tìm dòng ngày tháng trong toàn bộ tài liệu.
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        text = (para.text or '').strip().lower()
        if not text:
            continue
        if 'ngày' in text and 'tháng' in text and 'năm' in text:
            for run in para.runs:
                if run.text.strip():
                    set_font(run, bold=bool(run.bold), size=13, style_config=style_config)
                    run.font.italic = True
            return


# ---------------------------------------------------------------------------
# API chính
# ---------------------------------------------------------------------------

def render_document(
    template_path: Path,
    output_path: Path,
    metadata: Dict,
    blocks: List[Dict],
    style_config_path: Optional[Path] = None,
) -> None:
    """
    Render structured blocks vào TEMPLATE.docx theo thể thức KẾ HOẠCH.

    Args:
        template_path: File TEMPLATE.docx có placeholder {{ }}.
        output_path:   File .docx đầu ra.
        metadata:      Dict chứa các placeholder:
                           loai_van_ban, so_ky_hieu, ngay_thang, trich_yeu,
                           noi_nhan, nguoi_ky, chuc_vu_ky
        blocks:        List structured blocks từ parse_content_to_blocks().
        style_config_path: Đường dẫn style_config.json (mặc định dùng file cùng thư mục).
    """
    # --- Validation ---
    validation_errors = validate_blocks(blocks)
    if validation_errors:
        raise ValueError("Block validation failed:\n" + "\n".join(validation_errors))

    style_config = load_style_config(style_config_path)

    # --- Bước 1: Render metadata bằng docxtpl ---
    doc_tpl = DocxTemplate(template_path)
    data = {
        'loai_van_ban': metadata.get('loai_van_ban', 'KẾ HOẠCH'),
        'so_ky_hieu':   metadata.get('so_ky_hieu', ''),
        'ngay_thang':   metadata.get('ngay_thang', ''),
        'trich_yeu':    metadata.get('trich_yeu', ''),
        'noi_dung':     '',   # anchor - sẽ được thay bằng blocks ở bước 2
        'noi_nhan':     metadata.get('noi_nhan', ''),
        'nguoi_ky':     metadata.get('nguoi_ky', ''),
        'chuc_vu_ky':   metadata.get('chuc_vu_ky', ''),
    }
    doc_tpl.render(data)

    temp_path = output_path.parent / f"temp_{output_path.name}"
    doc_tpl.save(temp_path)

    # --- Bước 2: Chèn structured blocks vào vị trí anchor {{noi_dung}} ---
    doc = Document(temp_path)

    anchor_index = -1
    for i, para in enumerate(doc.paragraphs):
        if '{{noi_dung}}' in para.text.strip():
            anchor_index = i
            break

    if anchor_index == -1:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == '':
                anchor_index = i
                break

    if anchor_index == -1:
        print("⚠️  Không tìm thấy anchor {{noi_dung}}")
        anchor_index = len(doc.paragraphs) - 1

    body = doc._body._body
    body_children = list(body.iterchildren())
    anchor_element = doc.paragraphs[anchor_index]._element
    anchor_body_index = body_children.index(anchor_element)

    trailing_elements = [
        el for el in body_children[anchor_body_index + 1:]
        if not el.tag.endswith('sectPr')
    ]

    for element in trailing_elements:
        element.getparent().remove(element)
    anchor_element.getparent().remove(anchor_element)

    # Đường kẻ ngắn dưới tiêu đề (thể thức KẾ HOẠCH).
    _add_title_separator_line(doc, style_config)

    # Khoảng cách 1 dòng trống giữa tiêu đề và nội dung.
    if doc.paragraphs and doc.paragraphs[-1].text.strip() != '':
        doc.add_paragraph('')

    # Chèn blocks.
    styles = style_config.get('styles', {})
    for block in blocks:
        block_type = block['type']
        block_style = styles.get(block_type, styles.get('paragraph', {}))
        p = doc.add_paragraph(block['text'])
        p.clear()
        run = p.add_run(block['text'])
        set_paragraph_format(p, block_type, style_config)
        set_font(run, bold=block_style.get('bold', False), size=14, style_config=style_config)

    # Bỏ dòng trống dư ở đầu trailing_elements trước khi ghép lại.
    while trailing_elements and _is_empty_paragraph_element(trailing_elements[0]):
        trailing_elements.pop(0)

    if trailing_elements:
        if not doc.paragraphs or doc.paragraphs[-1].text.strip() != '':
            doc.add_paragraph('')

    # Ghép lại phần cuối tài liệu (bảng Nơi nhận / Ký tên).
    sect_pr = next(
        (el for el in body.iterchildren() if el.tag.endswith('sectPr')), None
    )
    for element in trailing_elements:
        if sect_pr is not None:
            body.insert(body.index(sect_pr), element)
        else:
            body.append(element)

    # Chuẩn hóa khoảng cách tiêu đề phụ.
    non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty_paras) >= 2:
        non_empty_paras[1].paragraph_format.space_before = Pt(0)

    # Post-processing chuẩn thể thức.
    _format_noi_nhan_table(doc, metadata.get('noi_nhan', ''), style_config)
    _format_date_line(doc, style_config)
    _normalize_default_header_blank_lines(doc)

    doc.save(output_path)
    temp_path.unlink()

    print(f"✅ Đã render {len(blocks)} blocks vào: {output_path}")


# ---------------------------------------------------------------------------
# Test / chạy trực tiếp
# ---------------------------------------------------------------------------

def test():
    """Test render KẾ HOẠCH với nội dung mẫu."""
    template_path = DEFAULT_TEMPLATE
    content_file = Path(__file__).parent / "noi_dung_ke_hoach.txt"
    output_path = Path(__file__).parent.parent / "van-ban-di" / "test_ke_hoach.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    content = content_file.read_text(encoding='utf-8')
    content = clean_content(content)
    blocks = parse_content_to_blocks(content)

    metadata = {
        'loai_van_ban': 'KẾ HOẠCH',
        'so_ky_hieu':   '123/KH-THPTĐBK',
        'ngay_thang':   'Đốc Bình Kiều, ngày 10 tháng 5 năm 2026',
        'trich_yeu':    'Triển khai thực hiện công tác bảo đảm trật tự, an toàn giao thông năm 2026',
        'noi_nhan':     'Sở GDĐT Đồng Tháp (báo cáo); Lưu: VT',
        'nguoi_ky':     'Nguyễn Minh Trí',
        'chuc_vu_ky':   'KT. HIỆU TRƯỞNG\nPHÓ HIỆU TRƯỞNG',
    }

    render_document(template_path, output_path, metadata, blocks)

    print(f"\n=== PARSED {len(blocks)} BLOCKS ===")
    for i, block in enumerate(blocks[:10]):
        print(f"{i:2d}: [{block['type']:12s}] {block['text'][:60]}...")


if __name__ == "__main__":
    test()
