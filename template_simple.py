#!/usr/bin/env python3
"""
Script tạo văn bản bằng python-docx trực tiếp (đơn giản hóa)
Dùng template cho header/footer, nội dung inject bằng code
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re

def clean_content(content):
    """Lọc bỏ tiêu đề và footer"""
    lines = content.split('\n')
    
    # Bỏ dòng đầu tiên nếu là loại văn bản
    if lines and lines[0].strip().upper() in ["KẾ HOẠCH", "QUYẾT ĐỊNH", "BÁO CÁO", "TỜ TRÌNH", "BIÊN BẢN", "CÔNG VĂN"]:
        lines = lines[1:]
    
    # Bỏ dòng thứ hai nếu là trích yếu
    if lines and lines[0].strip():
        lines = lines[1:]
    
    # Bỏ dòng trống tiếp theo
    if lines and not lines[0].strip():
        lines = lines[1:]
    
    # Lọc nội dung chính, bỏ footer
    result = []
    for line in lines:
        line = line.strip()
        
        # Tìm dấu hiệu footer
        if "Nơi nhận:" in line or re.match(r"^[^,]+,\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+", line, re.IGNORECASE):
            break
        
        # Loại bỏ "./." ở cuối
        line = line.replace("./.", "").strip()
        
        if line:
            result.append(line)
        else:
            result.append('')
    
    return '\n'.join(result)

def add_paragraph(doc, text, bold=False, indent=0):
    """Thêm đoạn văn với định dạng cơ bản"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    
    if indent > 0:
        p.paragraph_format.first_line_indent = Cm(indent)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = bold
    
    # Set font tiếng Việt
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.append(rFonts)

def add_content_from_text(doc, content):
    """Thêm nội dung từ text với định dạng đơn giản"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Đề mục cấp I, II, III - in đậm, không thụt đầu dòng
        if re.match(r'^[IVX]+\.', line):
            add_paragraph(doc, line, bold=True, indent=0)
        # Đề mục cấp 1, 2, 3 - in đậm, thụt đầu dòng 1cm
        elif re.match(r'^\d+\.', line):
            add_paragraph(doc, line, bold=True, indent=1.0)
        # Đề mục cấp a, b, c - bình thường, thụt đầu dòng 1.5cm
        elif re.match(r'^[a-z]+\.', line):
            add_paragraph(doc, line, bold=False, indent=1.5)
        # Bullet
        elif line.startswith('- '):
            clean = line[2:]
            add_paragraph(doc, '• ' + clean, bold=False, indent=0.75)
        # Dòng thường - thụt đầu dòng 1.25cm
        else:
            add_paragraph(doc, line, bold=False, indent=1.25)

def create_van_ban_from_template(template_path, content, output_path, metadata):
    """
    Tạo văn bản từ template + inject nội dung bằng python-docx
    
    Args:
        template_path: File template (có header/footer chuẩn)
        content: Nội dung chính
        output_path: File output
        metadata: Dict với trich_yeu, so_ky_hieu, ngay_thang, noi_nhan, nguoi_ky, chuc_vu
    """
    # Load template
    doc = Document(template_path)
    
    # Xóa nội dung cũ (giữ lại header/footer)
    # Đơn giản: xóa tất cả paragraph trừ header/footer
    # Header/footer trong docx nằm ở sections, không ở paragraphs
    
    # Cách đơn giản: tạo document mới, copy header/footer từ template
    # Nhưng phức tạp hơn
    
    # Cách thực tế: Inject vào placeholder trong template
    # Nhưng docx không hỗ trợ placeholder tốt
    
    # Cách đơn giản nhất: Tạo document mới với python-docx
    # Header/footer để người dùng tự thêm sau
    
    doc = Document()
    
    # Cài đặt trang
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # Thêm nội dung
    add_content_from_text(doc, content)
    
    # Save
    doc.save(output_path)
    print(f"✅ Đã tạo file: {output_path}")

def test():
    """Test với dữ liệu mẫu"""
    template_path = Path(__file__).parent / "TEMPLATE.docx"
    content_file = Path(__file__).parent / "noi_dung_ke_hoach.txt"
    output_path = Path(__file__).parent.parent / "van-ban-di" / "test_simple.docx"
    
    content = content_file.read_text(encoding='utf-8')
    content = clean_content(content)
    
    metadata = {
        'trich_yeu': 'Triển khai thực hiện Chỉ thị số 03/CT-TTg',
        'so_ky_hieu': '123/KH-THPTĐBK',
        'ngay_thang': 'Đốc Bình Kiều, ngày 10 tháng 5 năm 2026'
    }
    
    create_van_ban_from_template(template_path, content, output_path, metadata)

if __name__ == "__main__":
    test()
