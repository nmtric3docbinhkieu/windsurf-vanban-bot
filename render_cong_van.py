#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Render công văn (memo) từ TEMPLATE_CV.docx.
"""

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate

from renderer_engine import (
    _format_noi_nhan_table,
    _normalize_default_header_blank_lines,
    load_style_config,
    parse_content_to_blocks,
    set_font,
    set_paragraph_format,
    validate_blocks,
)


ROOT_DIR = Path(__file__).resolve().parent.parent
DEFAULT_INPUT_JSON = Path(__file__).resolve().parent / "noi_dung_cong_van.json"
DEFAULT_TEMPLATE = Path(__file__).resolve().parent / "TEMPLATE_CV.docx"
DEFAULT_OUTPUT_DIR = ROOT_DIR / "van-ban-di"


def _clear_cell(cell):
    cell.text = ""


def _set_cell_lines(cell, lines, font_sizes=None, default_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    _clear_cell(cell)
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.first_line_indent = Pt(0)
        size = font_sizes[index] if font_sizes and index < len(font_sizes) else default_size
        run = paragraph.add_run(line)
        set_font(run, bold=False, size=size)
        run.font.italic = italic


def _set_table_col_widths(table, widths_cm):
    """Set exact column widths for each cell in the table, and set table total width."""
    total_twips = sum(int(w * 567) for w in widths_cm)
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)
    tbl_w_el = OxmlElement("w:tblW")
    tbl_w_el.set(qn("w:w"), str(total_twips))
    tbl_w_el.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w_el)

    for row in table.rows:
        for cell_index, cell in enumerate(row.cells):
            if cell_index >= len(widths_cm):
                continue
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            for existing in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(existing)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(int(widths_cm[cell_index] * 567)))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)


def _find_table_with_label(doc, label):
    for table in doc.tables:
        try:
            if label in (table.cell(0, 0).text or ""):
                return table
        except Exception:
            continue
    return None


def _render_kinh_gui_table(doc, recipients):
    table = _find_table_with_label(doc, "Kính gửi")
    if table is None:
        return

    while len(table.rows) < max(3, len(recipients) + 1):
        table.add_row()

    # "Kính gửi:" canh phải, cột 1 nhỏ (1/3), cột 2 lớn (2/3)
    _set_cell_lines(table.cell(0, 0), ["Kính gửi:"], default_size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _clear_cell(table.cell(0, 1))

    for row_index, recipient in enumerate(recipients, start=1):
        _clear_cell(table.cell(row_index, 0))
        _set_cell_lines(table.cell(row_index, 1), [recipient], default_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    for row_index in range(len(recipients) + 1, len(table.rows)):
        _clear_cell(table.cell(row_index, 0))
        _clear_cell(table.cell(row_index, 1))

    # Cột 1 = 1/2 cột 2 → tỉ lệ 1:2, tổng 16cm (bằng vùng text A4 lề 2.5cm)
    _set_table_col_widths(table, [round(16/3, 2), round(32/3, 2)])

    # Đảm bảo spacing = 0 cho mọi paragraph trong mọi cell
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)


_ANCHOR_MARKER = "__ANCHOR_NOI_DUNG_CV__"


def _add_ky_ten_name(doc, name="Nguyễn Minh Trí"):
    """Thêm tên người ký vào cuối ô ký tên (right cell của bảng cuối)."""
    last_table = doc.tables[-1]
    right_cell = last_table.cell(0, 1)
    para = right_cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(name)
    set_font(run, bold=True, size=14)


def _remove_vml_lines(doc):
    """Xóa các đoạn chứa đường kẻ VML (w:pict) ra khỏi tài liệu."""
    body = doc._body._body
    to_remove = [
        el for el in body.findall(qn("w:p"))
        if el.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict") is not None
    ]
    for el in to_remove:
        el.getparent().remove(el)


def _remove_extra_blank_after_element(doc, target_element, keep=1):
    """Giảm blank paragraphs ngay SAU target_element xuống còn keep."""
    body = doc._body._body
    children = list(body.iterchildren())
    try:
        idx = children.index(target_element)
    except ValueError:
        return
    blank_paras = []
    for el in children[idx + 1:]:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            text = "".join(r.text or "" for r in el.iter(qn("w:t"))).strip()
            if text == "":
                blank_paras.append(el)
            else:
                break
        else:
            break
    # remove from beginning (furthest from next table, closest to content above)
    excess = len(blank_paras) - keep
    for el in blank_paras[:excess]:
        el.getparent().remove(el)


def _remove_extra_blank_before_element(doc, target_element, keep=1):
    """Giảm blank paragraphs ngay TRƯỚC target_element xuống còn keep."""
    body = doc._body._body
    children = list(body.iterchildren())
    try:
        idx = children.index(target_element)
    except ValueError:
        return
    blank_paras = []
    for el in reversed(children[:idx]):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            text = "".join(r.text or "" for r in el.iter(qn("w:t"))).strip()
            if text == "":
                blank_paras.append(el)
            else:
                break
        else:
            break
    # blank_paras[0] = closest to target, blank_paras[-1] = furthest
    # keep the ones closest to target, remove the rest
    excess = len(blank_paras) - keep
    for el in blank_paras[keep:]:
        el.getparent().remove(el)


def _insert_blocks(doc, blocks):
    anchor_index = -1
    for index, paragraph in enumerate(doc.paragraphs):
        if _ANCHOR_MARKER in (paragraph.text or ""):
            anchor_index = index
            break

    if anchor_index == -1:
        # fallback nếu marker không tìm thấy
        for index, paragraph in enumerate(doc.paragraphs):
            if (paragraph.text or "").strip() == "":
                anchor_index = index
                break

    if anchor_index == -1:
        anchor_index = len(doc.paragraphs) - 1

    body = doc._body._body
    body_children = list(body.iterchildren())
    anchor_element = doc.paragraphs[anchor_index]._element
    anchor_body_index = body_children.index(anchor_element)

    trailing_elements = []
    for element in body_children[anchor_body_index + 1:]:
        if element.tag.endswith("sectPr"):
            continue
        trailing_elements.append(element)

    for element in trailing_elements:
        element.getparent().remove(element)
    anchor_element.getparent().remove(anchor_element)

    style_config = load_style_config()
    styles = style_config.get("styles", {})

    for block in blocks:
        block_type = block["type"]
        block_style = styles.get(block_type, styles.get("paragraph", {}))
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(block["text"])
        set_paragraph_format(paragraph, block_type, style_config)
        set_font(run, bold=block_style.get("bold", False), size=14, style_config=style_config)

    sect_pr = None
    for element in body.iterchildren():
        if element.tag.endswith("sectPr"):
            sect_pr = element
            break

    if trailing_elements and (not doc.paragraphs or doc.paragraphs[-1].text.strip() != ""):
        doc.add_paragraph("")

    for element in trailing_elements:
        if sect_pr is not None:
            body.insert(body.index(sect_pr), element)
        else:
            body.append(element)


def _format_cong_van_header(doc, so_ky_hieu, trich_yeu):
    if not doc.tables:
        return

    header_table = doc.tables[0]
    if len(header_table.rows) < 2 or len(header_table.columns) < 2:
        return

    # Số: size 13, trích yếu: size 12, canh giữa
    _set_cell_lines(
        header_table.cell(1, 0),
        [f"Số: {so_ky_hieu}", trich_yeu],
        font_sizes=[13, 12],
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Ngày tháng: size 13, in nghiêng, có chữ "ngày"
    now = datetime.now()
    date_str = f"Đồng Tháp, ngày {now.day} tháng {now.month} năm {now.year}"
    _set_cell_lines(
        header_table.cell(1, 1),
        [date_str],
        font_sizes=[13],
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Xóa hàng dư thứ 3 trong header table
    if len(header_table.rows) > 2:
        row_element = header_table.rows[2]._tr
        row_element.getparent().remove(row_element)


def _load_content_source(noi_dung_path):
    with open(noi_dung_path, "r", encoding="utf-8") as f:
        return json.load(f)


def render_cong_van(noi_dung_path=DEFAULT_INPUT_JSON, output_dir=DEFAULT_OUTPUT_DIR, timestamp=None):
    """Render công văn từ JSON content."""

    noi_dung = _load_content_source(noi_dung_path)

    required_fields = ["loai_van_ban", "so_ky_hieu_goi_y", "trich_yeu", "noi_dung"]
    for field in required_fields:
        if field not in noi_dung:
            raise ValueError(f"Missing required field: {field}")

    if noi_dung.get("loai_van_ban") != "cong_van":
        raise ValueError(f"Expected loai_van_ban='cong_van', got '{noi_dung.get('loai_van_ban')}'")

    blocks = parse_content_to_blocks(noi_dung.get("noi_dung", ""))
    errors = validate_blocks(blocks)
    if errors:
        raise ValueError("Block validation failed:\n" + "\n".join(errors))

    if timestamp is None:
        timestamp = datetime.now().strftime("%Y_%m_%d_%H%M%S")

    so_ky_hieu = noi_dung.get("so_ky_hieu_goi_y", "").replace("/", "_").replace(" ", "").strip("_")
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"VBDi_{so_ky_hieu}_CV_{timestamp}.docx"

    template_path = Path(DEFAULT_TEMPLATE)
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    metadata = {
        "so_ky_hieu": noi_dung.get("so_ky_hieu_goi_y", ""),
        "trich_yeu": noi_dung.get("trich_yeu", ""),
        "noi_nhan": noi_dung.get("noi_nhan", ""),
        "kinh_gui": noi_dung.get("kinh_gui", []),
        "ngay_thang": f"Ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}",
    }

    doc_tpl = DocxTemplate(template_path)
    doc_tpl.render({
        "so_ky_hieu": metadata["so_ky_hieu"],
        "trich_yeu": metadata["trich_yeu"],
        "noi_nhan": metadata["noi_nhan"],
        "kinh_gui": metadata["kinh_gui"],
        "ngay_thang": metadata["ngay_thang"],
        "noi_dung": _ANCHOR_MARKER,  # unique marker để _insert_blocks tìm đúng anchor
    })

    temp_path = output_dir / f"temp_{output_path.name}"
    doc_tpl.save(temp_path)

    doc = Document(temp_path)
    _insert_blocks(doc, blocks)
    _format_cong_van_header(doc, metadata["so_ky_hieu"], metadata["trich_yeu"])
    _render_kinh_gui_table(doc, metadata["kinh_gui"])
    _format_noi_nhan_table(doc, metadata["noi_nhan"])
    _add_ky_ten_name(doc)
    _normalize_default_header_blank_lines(doc)
    _remove_vml_lines(doc)
    _remove_extra_blank_after_element(doc, doc.tables[0]._tbl, keep=1)
    _remove_extra_blank_before_element(doc, doc.tables[-1]._tbl, keep=1)
    doc.save(output_path)
    temp_path.unlink(missing_ok=True)

    print(f"✅ Đã render {len(blocks)} blocks vào: {output_path}")
    return output_path


if __name__ == "__main__":
    try:
        output_path = render_cong_van()
        print(f"\n✅ SUCCESS: {output_path}")
    except Exception as exc:
        print(f"\n❌ ERROR: {exc}")
        raise