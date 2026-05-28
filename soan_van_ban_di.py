"""
Module: soan_van_ban_di.py
Mục đích: Đọc văn bản đến từ Sở GDĐT, dùng AI soạn văn bản đi phù hợp
với đặc thù Trường THPT Đốc Binh Kiều.
"""

import os
import re
import json
import subprocess
from pathlib import Path
from datetime import datetime
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

# Load biến môi trường từ file .env
load_dotenv()

# ============================================================
# CẤU HÌNH
# ============================================================
THONG_TIN_TRUONG = """
- Tên đầy đủ: Trường Trung học Phổ thông Đốc Binh Kiều
- Địa chỉ: Ấp 3, xã Đốc Binh Kiều, huyện Tháp Mười, tỉnh Đồng Tháp
- Điện thoại: 02773826878 | Email: thptdocbinhkieu@dongthap.edu.vn
- Cơ quan chủ quản: Sở Giáo dục và Đào tạo Đồng Tháp
- Người ký văn bản: PHÓ HIỆU TRƯỞNG Nguyễn Minh Trí (đang khuyết Hiệu trưởng)
- Tổng số lớp: 13 lớp (Khối 10: 4 lớp, Khối 11: 5 lớp, Khối 12: 4 lớp)
- Tổng số học sinh: 506 em (Nam: 274, Nữ: 232)
- Tổng số giáo viên: 30 giáo viên
- Đặc điểm: Trường vùng nông thôn, học sinh chủ yếu từ xã Đốc Binh Kiều,
  xã Tháp Mười và một số xã thuộc tỉnh Tây Ninh giáp ranh.
  Điều kiện kinh tế gia đình học sinh còn nhiều khó khăn.
  Giáo viên nhiệt tình, đoàn kết; học sinh chăm ngoan, có ý thức học tập.
- Thành tích nổi bật: Trường đạt chuẩn quốc gia, hoàn thành tốt các chỉ tiêu
  chuyên môn hằng năm, có học sinh đạt giải thi học sinh giỏi cấp tỉnh.
- Tên viết tắt dùng trong văn bản: Trường THPT Đốc Binh Kiều
"""

# Tiêu đề cơ quan ban hành (góc trái)
CO_QUAN_BAN_HANH = "TRƯỜNG THPT ĐỐC BÌNH KIỀU"
CO_QUAN_CHU_QUAN = "SỞ GDĐT ĐỒNG THÁP"

# Thư mục - điều chỉnh đúng với thực tế
THU_MUC_VAN_BAN_DEN = "../van-ban-den-xu-ly"  # thư mục chứa văn bản đến đã xử lý
THU_MUC_VAN_BAN_DI  = "../van-ban-di"          # thư mục xuất văn bản đi

# DeepSeek API Configuration
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEEPSEEK_MODEL = "deepseek-chat"  # Model phổ biến, giá hợp lý, đã test hoạt động


def _env_flag(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


# Mặc định TẮT gọi API để tiết kiệm chi phí khi phát triển.
# Bật lại bằng cách set USE_DEEPSEEK_API=1 khi triển khai chính thức.
USE_DEEPSEEK_API = _env_flag("USE_DEEPSEEK_API", False)

ALLOWED_DOC_TYPES = {
    "cong_van",
    "ke_hoach",
    "bao_cao",
    "to_trinh",
    "quyet_dinh",
    "bien_ban",
}

SO_KY_HIEU_BY_TYPE = {
    "cong_van": "/THPTĐBK",
    "ke_hoach": "/KH-THPTĐBK",
    "bao_cao": "/BC-THPTĐBK",
    "to_trinh": "/TTr-THPTĐBK",
    "quyet_dinh": "/QĐ-THPTĐBK",
    "bien_ban": "/BB-THPTĐBK",
}

REQUIRED_JSON_KEYS = [
    "loai_van_ban",
    "so_ky_hieu_goi_y",
    "trich_yeu",
    "tom_tat_yeu_cau",
    "noi_nhan",
    "noi_dung",
]

HO_SO_TRUONG_FILE = "THONG TIN TRUONG THPT DOC BINH KIEU 2026.docx"
MAX_PROFILE_CHARS = 12000
MAX_SOURCE_CHARS = 14000
MAX_STRUCTURE_LINES = 80
MAX_SOURCE_PREVIEW_LINES = 220

_SCHOOL_CONTEXT_CACHE = None


def dinh_dang_ngay_thang_hien_tai() -> str:
    """Trả về chuỗi ngày tháng để gắn vào template.

    Template đã chứa sẵn tiền tố "Đồng Tháp, ", nên ở đây chỉ trả về phần
    "ngày ... tháng ... năm ...".
    """
    hom_nay = datetime.now()
    return f"ngày {hom_nay.day} tháng {hom_nay.month} năm {hom_nay.year}"


def dinh_dang_noi_nhan(noi_nhan_raw: str, noi_dung: str = "") -> str:
    """Định dạng nơi nhận theo thể thức gạch đầu dòng như mẫu Sở.

    Ví dụ input:
        "Sở GDĐT Đồng Tháp (báo cáo); Lưu: VT"
    Output:
        "- Sở GDĐT Đồng Tháp (báo cáo);\n- Lưu: VT."
    """
    if not noi_nhan_raw:
        noi_nhan_raw = ""

    parts = [p.strip() for p in re.split(r';\s*', noi_nhan_raw) if p.strip()]

    # Nhóm người nhận bắt buộc theo yêu cầu nghiệp vụ hiện tại
    bat_buoc = ["Phó Hiệu trưởng", "Tổ trưởng", "Giáo viên chủ nhiệm"]

    # Tự động bổ sung nơi nhận theo bộ phận được nhắc trong nội dung
    noi_dung_lower = (noi_dung or "").lower()
    goi_y_theo_noi_dung = [
        ("đoàn trường", "Đoàn trường"),
        ("giáo viên chủ nhiệm", "Giáo viên chủ nhiệm"),
        ("tổ trưởng", "Tổ trưởng"),
        ("phó hiệu trưởng", "Phó Hiệu trưởng"),
    ]

    # Loại bỏ nơi nhận không dùng theo yêu cầu nghiệp vụ hiện tại.
    parts_filtered = []
    for part in parts:
        p_lower = part.lower()
        if "ban giám hiệu" in p_lower:
            continue
        if "văn phòng" in p_lower or "van phong" in p_lower:
            continue
        if "lưu" in p_lower and "vt" in p_lower:
            continue
        parts_filtered.append(part)
    parts = parts_filtered

    lower_parts = " | ".join(parts).lower()
    for muc in bat_buoc:
        if muc.lower() not in lower_parts:
            parts.append(muc)
            lower_parts += f" | {muc.lower()}"

    for keyword, muc in goi_y_theo_noi_dung:
        if keyword in noi_dung_lower and muc.lower() not in lower_parts:
            parts.append(muc)
            lower_parts += f" | {muc.lower()}"

    # Luôn đặt "Lưu: VT." ở cuối danh sách nơi nhận.
    parts = [p for p in parts if not ("lưu" in p.lower() and "vt" in p.lower())]
    parts.append("Lưu: VT")

    if not parts:
        return ""

    lines = []
    for idx, part in enumerate(parts):
        clean_part = part.rstrip(';').rstrip('.')
        line_end = '.' if idx == len(parts) - 1 else ';'
        lines.append(f"- {clean_part}{line_end}")

    return "\n".join(lines)

# ============================================================
# BƯỚC 1: ĐỌC NỘI DUNG FILE VĂN BẢN ĐẾN
# ============================================================

def doc_noi_dung_file(duong_dan: str) -> str:
    """Đọc nội dung văn bản đến (docx hoặc pdf)."""
    ext = Path(duong_dan).suffix.lower()
    
    if ext == ".docx":
        try:
            # Ưu tiên đọc đầy đủ cả paragraph + bảng để không mất cấu trúc văn bản nguồn.
            noi_dung_day_du = _doc_docx_day_du(Path(duong_dan))
            if noi_dung_day_du.strip():
                return noi_dung_day_du

            # Fallback nếu file có định dạng đặc biệt.
            doc = Document(duong_dan)
            noi_dung = []
            for para in doc.paragraphs:
                if para.text.strip():
                    noi_dung.append(para.text.strip())
            return "\n".join(noi_dung)
        except Exception as e:
            print(f"⚠️  Lỗi đọc file DOCX: {e}")
            return f"[Không đọc được file DOCX: {e}]"
    
    elif ext == ".pdf":
        # Gọi script Node.js pdf-reader.js để đọc PDF
        try:
            script_path = Path(__file__).parent / "pdf-reader.js"
            result = subprocess.run(
                ["node", "-e", f"""
const readPDF = require('{script_path}').readPDF;
readPDF('{duong_dan}').then(result => {{
    console.log(result.text);
}}).catch(err => {{
    console.error('ERROR:', err.message);
    process.exit(1);
}});
"""],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
            else:
                return f"[Không đọc được file PDF: {result.stderr}]"
        except Exception as e:
            return f"[Không đọc được file PDF: {e}]"
    
    else:
        return f"[Định dạng file không hỗ trợ: {ext}]"


# ============================================================
# BƯỚC 2: PHÂN TÍCH LOẠI VĂN BẢN VÀ GỌI AI SOẠN
# ============================================================

def _doc_docx_day_du(file_path: Path) -> str:
    """Đọc cả paragraph và bảng từ file DOCX để lấy tối đa thông tin ngữ cảnh."""
    doc = Document(file_path)
    chunks = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = " ".join(
                    p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()
                ).strip()
                if text:
                    cells.append(text)
            if cells:
                chunks.append(" | ".join(cells))

    raw = "\n".join(chunks)
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw).strip()
    return raw


def _tai_ho_so_truong() -> str:
    """Ưu tiên đọc hồ sơ trường từ file mới; fallback về thông tin tĩnh nếu thiếu file."""
    profile_path = Path(__file__).parent / HO_SO_TRUONG_FILE
    if not profile_path.exists():
        print(f"⚠️  Không tìm thấy hồ sơ trường: {profile_path}. Dùng thông tin tĩnh.")
        return THONG_TIN_TRUONG.strip()

    try:
        profile_text = _doc_docx_day_du(profile_path)
        if not profile_text:
            return THONG_TIN_TRUONG.strip()

        if len(profile_text) > MAX_PROFILE_CHARS:
            profile_text = profile_text[:MAX_PROFILE_CHARS].rstrip() + "\n[...đã rút gọn để tiết kiệm token...]"

        return profile_text
    except Exception as exc:
        print(f"⚠️  Không đọc được hồ sơ trường ({exc}). Dùng thông tin tĩnh.")
        return THONG_TIN_TRUONG.strip()


def _parse_int(value: str) -> str:
    digits = re.sub(r"\D+", "", value or "")
    return digits if digits else ""


def _tach_cac_dong_dac_diem(profile_text: str, max_items: int = 6) -> list:
    """Lấy một số dòng đặc điểm tiêu biểu để AI tham chiếu nhanh."""
    lines = [ln.strip(" -\t") for ln in (profile_text or "").splitlines() if ln.strip()]
    patterns = [
        "đặc điểm",
        "thuận lợi",
        "khó khăn",
        "mục tiêu",
        "định hướng",
        "cơ sở vật chất",
        "đội ngũ",
        "học sinh",
        "kết quả",
    ]
    picked = []
    for line in lines:
        lower = line.lower()
        if re.match(r"^[IVXLCDM]+\.", line):
            continue
        if line.isupper() and len(line) <= 40:
            continue
        if any(p in lower for p in patterns):
            line = line[:220].rstrip() + "..." if len(line) > 220 else line
            picked.append(line)
        if len(picked) >= max_items:
            break
    return picked


def _trich_thong_tin_truong_quan_trong(profile_text: str) -> dict:
    """Trích các dữ liệu trọng yếu thành biến riêng, giúp AI dùng ổn định."""
    text = profile_text or ""

    def find(pattern: str, default: str = "") -> str:
        m = re.search(pattern, text, flags=re.IGNORECASE)
        return m.group(1).strip() if m else default

    ten_truong = find(r"(?:Tên trường|Tên đầy đủ)\s*:\s*([^\n]+)", "Trường THPT Đốc Binh Kiều")
    dia_chi = find(r"Địa chỉ\s*:\s*([^\n]+)")
    co_quan_chu_quan = find(r"Cơ quan chủ quản\s*:\s*([^\n]+)", "Sở Giáo dục và Đào tạo Đồng Tháp")
    dien_thoai = find(r"Điện thoại\s*:\s*([0-9+().\s-]+)")
    if not dien_thoai:
        phone_match = re.search(r"(0\d{8,10})", text)
        dien_thoai = phone_match.group(1) if phone_match else ""
    email = find(r"Email\s*:\s*([^\s\n]+)")
    website = find(r"Website\s*:\s*([^\s\n]+)")
    nam_thanh_lap = find(r"Năm thành lập\s*:\s*([^\n]+)")

    tong_lop = _parse_int(find(r"(?:Tổng số lớp|Số lớp)\s*:\s*([^\n]+)")) or "13"
    tong_hoc_sinh = _parse_int(find(r"(?:Tổng số học sinh|Số học sinh)\s*:\s*([^\n]+)")) or "506"
    tong_giao_vien = _parse_int(find(r"(?:Tổng số giáo viên|Số giáo viên)\s*:\s*([^\n]+)")) or "30"
    hoc_sinh_nam = _parse_int(find(r"Nam\s*:\s*([^,;\n]+)"))
    hoc_sinh_nu = _parse_int(find(r"Nữ\s*:\s*([^,;\n]+)"))

    dac_diem = _tach_cac_dong_dac_diem(text, max_items=4)

    return {
        "ten_truong": ten_truong,
        "dia_chi": dia_chi,
        "co_quan_chu_quan": co_quan_chu_quan,
        "dien_thoai": dien_thoai,
        "email": email,
        "website": website,
        "nam_thanh_lap": nam_thanh_lap,
        "tong_lop": tong_lop,
        "tong_hoc_sinh": tong_hoc_sinh,
        "tong_giao_vien": tong_giao_vien,
        "hoc_sinh_nam": hoc_sinh_nam,
        "hoc_sinh_nu": hoc_sinh_nu,
        "dac_diem_noi_bat": dac_diem,
    }


def _xay_dung_context_truong_toi_gian(profile_text: str) -> str:
    """Tạo context cực gọn để giảm token nhưng vẫn đủ dữ liệu cốt lõi."""
    facts = _trich_thong_tin_truong_quan_trong(profile_text)
    items = [
        f"ten_truong={facts['ten_truong']}",
        f"co_quan_chu_quan={facts['co_quan_chu_quan']}",
        f"dia_chi={facts['dia_chi']}",
        f"tong_lop={facts['tong_lop']}",
        f"tong_hoc_sinh={facts['tong_hoc_sinh']}",
        f"tong_giao_vien={facts['tong_giao_vien']}",
    ]
    if facts["hoc_sinh_nam"] or facts["hoc_sinh_nu"]:
        items.append(f"hoc_sinh_nam={facts['hoc_sinh_nam']};hoc_sinh_nu={facts['hoc_sinh_nu']}")
    if facts["dien_thoai"]:
        items.append(f"dien_thoai={facts['dien_thoai']}")
    if facts["email"]:
        items.append(f"email={facts['email']}")
    if facts["website"]:
        items.append(f"website={facts['website']}")
    if facts["nam_thanh_lap"]:
        items.append(f"nam_thanh_lap={facts['nam_thanh_lap']}")

    if facts["dac_diem_noi_bat"]:
        items.append("dac_diem_noi_bat=" + " | ".join(facts["dac_diem_noi_bat"]))

    return "\n".join(items).strip()


def _lay_context_truong_toi_uu() -> str:
    """Nạp context trường 1 lần và tái sử dụng để ổn định + tiết kiệm token."""
    global _SCHOOL_CONTEXT_CACHE
    if _SCHOOL_CONTEXT_CACHE:
        return _SCHOOL_CONTEXT_CACHE

    profile_text = _tai_ho_so_truong()
    _SCHOOL_CONTEXT_CACHE = _xay_dung_context_truong_toi_gian(profile_text)
    return _SCHOOL_CONTEXT_CACHE


def _rut_ban_do_cau_truc_van_ban(content: str) -> str:
    """Rút các dòng tiêu đề/đánh số làm bản đồ cấu trúc để AI bám khung nguồn."""
    lines = [ln.strip() for ln in (content or "").splitlines() if ln.strip()]
    structure = []
    for line in lines:
        if re.match(r"^(PHẦN|CHƯƠNG|MỤC)\b", line, re.IGNORECASE):
            structure.append(line)
            continue
        if re.match(r"^[IVXLCDM]+\.", line):
            structure.append(line)
            continue
        if re.match(r"^\d+[\.)]", line):
            structure.append(line)
            continue
        if re.match(r"^[a-zA-Z][\.)]", line):
            structure.append(line)
            continue

    if not structure:
        structure = lines[:20]

    return "\n".join(structure[:MAX_STRUCTURE_LINES])


def _rut_gon_noi_dung_nguon(content: str) -> str:
    """Rút gọn nội dung nguồn: giữ thứ tự dòng, bỏ dòng rỗng dư để tiết kiệm token."""
    lines = [ln.strip() for ln in (content or "").splitlines() if ln.strip()]
    compact = "\n".join(lines[:MAX_SOURCE_PREVIEW_LINES])
    if len(compact) > MAX_SOURCE_CHARS:
        compact = compact[:MAX_SOURCE_CHARS].rstrip() + "\n[...đã rút gọn...]"
    return compact


def _parse_context_kv(context_text: str) -> dict:
    """Parse context key=value thành dict để dùng cho bộ soạn offline."""
    data = {}
    for line in (context_text or "").splitlines():
        line = line.strip()
        if not line or "=" not in line:
            continue
        key, value = line.split("=", 1)
        data[key.strip()] = value.strip()
    return data


def _xac_dinh_loai_van_ban_tu_nguon(content: str) -> str:
    text = (content or "").upper()
    if "KẾ HOẠCH" in text:
        return "ke_hoach"
    if "BÁO CÁO" in text:
        return "bao_cao"
    if "TỜ TRÌNH" in text:
        return "to_trinh"
    if "QUYẾT ĐỊNH" in text:
        return "quyet_dinh"
    if "BIÊN BẢN" in text:
        return "bien_ban"
    return "cong_van"


def _trich_trich_yeu_tu_nguon(content: str, loai_vb: str) -> str:
    lines = [ln.strip() for ln in (content or "").splitlines() if ln.strip()]
    if not lines:
        return "Triển khai nhiệm vụ theo văn bản chỉ đạo"

    title_tokens = {
        "KẾ HOẠCH", "BÁO CÁO", "TỜ TRÌNH", "QUYẾT ĐỊNH", "BIÊN BẢN", "CÔNG VĂN"
    }
    for line in lines[:12]:
        up = line.upper()
        if up in title_tokens:
            continue
        if len(line) > 12:
            return line.rstrip(".")

    if loai_vb == "ke_hoach":
        return "Tổ chức thực hiện nhiệm vụ theo chỉ đạo của Sở GDĐT"
    return "Thực hiện nhiệm vụ theo chỉ đạo của Sở GDĐT"


def _adapt_line_to_thpt(line: str, school_name: str) -> str:
    """Điều chỉnh câu chữ từ phạm vi Sở về phạm vi nhà trường nhưng giữ khung mục."""
    text = line
    replacements = [
        ("Sở Giáo dục và Đào tạo", school_name),
        ("Sở GDĐT", school_name),
        ("ngành Giáo dục", "nhà trường"),
        ("các phòng chuyên môn thuộc Sở Giáo dục và Đào tạo", "các bộ phận chuyên môn của nhà trường"),
        ("Các phòng chuyên môn thuộc Sở Giáo dục và Đào tạo", "Các bộ phận chuyên môn của nhà trường"),
        ("các đơn vị trực thuộc Sở Giáo dục và Đào tạo", "các tổ chuyên môn, Đoàn trường và bộ phận liên quan"),
        ("Các đơn vị trực thuộc Sở Giáo dục và Đào tạo", "Các tổ chuyên môn, Đoàn trường và bộ phận liên quan"),
        ("Thủ trưởng các cơ sở giáo dục", "Hiệu trưởng/Phó Hiệu trưởng phụ trách và các bộ phận liên quan"),
    ]
    for src, dst in replacements:
        text = text.replace(src, dst)
    return text


def _is_heading_line(line: str) -> bool:
    roman = re.match(r"^([IVXLCDM]+)\.\s+(.+)$", line)
    if roman:
        title = roman.group(2).strip()
        return len(title) <= 140

    numeric = re.match(r"^(\d+(?:\.\d+)*)\.\s+(.+)$", line)
    if numeric:
        title = numeric.group(2).strip()
        # Tiêu đề thường ngắn, không kết thúc bằng dấu chấm.
        return len(title) <= 120 and not title.endswith(".")

    if re.match(r"^[a-zA-Z]\)\s+", line):
        return True
    if re.match(r"^[a-zA-Z]\.\s+", line):
        return True
    if re.match(r"^(PHẦN|CHƯƠNG|MỤC)\b", line, re.IGNORECASE):
        return True
    return False


def _rewrite_intro_line(line: str, school_name: str) -> str:
    text = _adapt_line_to_thpt(line, school_name)
    if text.upper().startswith("CĂN CỨ"):
        return text
    return f"{school_name} xây dựng kế hoạch thực hiện cụ thể như sau:"


def _renumber_section_items(lines: list, section_prefix: str = "IV.") -> list:
    """Đánh số lại các đoạn trong một mục La Mã theo thứ tự 1..N."""
    if not lines:
        return lines

    start_idx = -1
    for i, line in enumerate(lines):
        if line.strip().startswith(section_prefix):
            start_idx = i
            break

    if start_idx == -1:
        return lines

    end_idx = len(lines)
    for i in range(start_idx + 1, len(lines)):
        if re.match(r"^[IVXLCDM]+\.\s+", lines[i].strip()):
            end_idx = i
            break

    result = list(lines)
    count = 1
    for i in range(start_idx + 1, end_idx):
        raw = result[i].strip()
        if not raw:
            continue
        if _is_heading_line(raw):
            continue

        content = re.sub(r"^\d+(?:\.\d+)*\.\s*", "", raw).strip()
        result[i] = f"{count}. {content}"
        count += 1

    return result


def _tao_can_cu_chinh_tu_van_ban(noi_dung_van_ban_den: str, loai_vb: str, trich_yeu: str) -> str:
    """Sinh 1 dòng căn cứ ngắn gọn, chỉ bám văn bản chính đang xử lý."""
    text = noi_dung_van_ban_den or ""
    loai_map = {
        "ke_hoach": "Kế hoạch",
        "cong_van": "Công văn",
        "bao_cao": "Báo cáo",
        "to_trinh": "Tờ trình",
        "quyet_dinh": "Quyết định",
        "bien_ban": "Biên bản",
    }
    loai_label = loai_map.get(loai_vb, "Văn bản")

    so_van_ban = ""
    date_text = ""
    so_span = None

    so_match = re.search(r"\b(?:Số|số)\s*:\s*([0-9]+/[A-ZĐa-z\-]+)", text)
    if not so_match:
        so_match = re.search(r"\b(?:Kế hoạch|Công văn|Báo cáo|Tờ trình|Quyết định|Biên bản)\s+số\s+([0-9]+/[A-ZĐa-z\-]+)", text, re.IGNORECASE)
    if so_match:
        so_van_ban = so_match.group(1).strip()
        so_span = so_match.span()

    if so_span:
        start = max(0, so_span[0] - 120)
        end = min(len(text), so_span[1] + 220)
        nearby = text[start:end]
        date_match = re.search(r"ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}", nearby, re.IGNORECASE)
        if date_match:
            date_text = date_match.group(0).strip()

    if not date_text:
        date_match = re.search(r"ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}", text, re.IGNORECASE)
        if date_match:
            date_text = date_match.group(0).strip()

    if so_van_ban and date_text:
        return f"Căn cứ {loai_label} số {so_van_ban} {date_text} của Sở Giáo dục và Đào tạo Đồng Tháp về {trich_yeu.lower()};"
    if so_van_ban:
        return f"Căn cứ {loai_label} số {so_van_ban} của Sở Giáo dục và Đào tạo Đồng Tháp về {trich_yeu.lower()};"
    return f"Căn cứ văn bản chỉ đạo của Sở Giáo dục và Đào tạo Đồng Tháp về {trich_yeu.lower()};"


def _rewrite_body_line(current_h1: str, current_h2: str, source_line: str, school_ctx: dict, idx: int) -> str:
    school_name = school_ctx.get("ten_truong", "Trường THPT Đốc Binh Kiều")
    tong_lop = school_ctx.get("tong_lop", "13")
    tong_hoc_sinh = school_ctx.get("tong_hoc_sinh", "506")
    tong_giao_vien = school_ctx.get("tong_giao_vien", "30")

    hint = f"{current_h1} {current_h2} {source_line}".lower()

    def pick(options: list) -> str:
        return options[idx % len(options)]

    if "mục đích" in hint or "yêu cầu" in hint:
        return pick([
            f"Tập trung nâng cao ý thức chấp hành pháp luật giao thông cho {tong_hoc_sinh} học sinh thông qua giáo dục thường xuyên, gắn với tiêu chí thi đua của tập thể lớp và cá nhân học sinh.",
            "Đặt mục tiêu giảm vi phạm giao thông trong học sinh bằng biện pháp giáo dục sớm, theo dõi thường xuyên và phối hợp chặt giữa giáo viên chủ nhiệm với phụ huynh.",
            "Bảo đảm các yêu cầu triển khai có tính khả thi, đo được kết quả và phù hợp điều kiện dạy học thực tế của nhà trường."
        ])
    if "tuyên truyền" in hint or "phổ biến" in hint or "giáo dục pháp luật" in hint:
        return pick([
            "Giao Đoàn trường phối hợp giáo viên chủ nhiệm tổ chức tuyên truyền theo từng khối lớp, lồng ghép trong sinh hoạt dưới cờ, sinh hoạt lớp và các kênh thông tin của nhà trường.",
            "Tăng cường truyền thông trực quan tại lớp học, khu vực cổng trường và nhóm liên lạc phụ huynh để nhắc nhở học sinh chấp hành an toàn giao thông hằng ngày.",
            "Mỗi học kỳ tổ chức chuyên đề tuyên truyền trọng tâm về các hành vi dễ vi phạm, kết hợp trao đổi tình huống thực tế gần trường và trên tuyến đường học sinh thường đi lại."
        ])
    if "cam kết" in hint or "mô hình an toàn" in hint or "cổng trường" in hint:
        return pick([
            "Tổ chức ký cam kết thực hiện an toàn giao thông đối với học sinh, giáo viên và phụ huynh; duy trì nền nếp cổng trường an toàn vào các khung giờ cao điểm.",
            "Rà soát việc đội mũ bảo hiểm, chấp hành tốc độ và điểm dừng đón trả học sinh; xử lý nhắc nhở theo quy định nội bộ của nhà trường.",
            "Phân công lực lượng trực cổng phối hợp bảo vệ, giáo viên trực tuần để giảm ùn tắc và phòng ngừa nguy cơ mất an toàn trước cổng trường."
        ])
    if "phối hợp" in hint or "gia đình" in hint or "xã hội" in hint or "công an" in hint:
        return pick([
            "Tăng cường phối hợp giữa nhà trường, cha mẹ học sinh và công an địa phương để nắm tình hình, kịp thời nhắc nhở, giáo dục các trường hợp có nguy cơ vi phạm.",
            "Duy trì kênh trao đổi hai chiều giữa giáo viên chủ nhiệm và phụ huynh để xử lý sớm các dấu hiệu vi phạm trật tự an toàn giao thông.",
            "Khi phát sinh vụ việc, nhà trường phối hợp cơ quan chức năng và gia đình để thống nhất biện pháp giáo dục, tránh tái phạm."
        ])
    if "phân công" in hint or "văn phòng" in hint or "tổ chuyên môn" in hint or "đơn vị trực thuộc" in hint:
        return pick([
            f"Phân công rõ trách nhiệm cho Ban giám hiệu, {tong_giao_vien} giáo viên và các bộ phận liên quan; định kỳ hằng tháng rà soát tiến độ thực hiện và xử lý vướng mắc phát sinh.",
            "Văn phòng theo dõi tiến độ, tổng hợp số liệu; các tổ chuyên môn và giáo viên chủ nhiệm chịu trách nhiệm triển khai trực tiếp đến từng lớp.",
            "Đoàn trường, tổ giám thị và giáo viên chủ nhiệm phối hợp kiểm tra nền nếp, báo cáo kết quả theo mốc thời gian đã thống nhất."
        ])
    if "tổ chức thực hiện" in hint or "kiểm tra" in hint or "báo cáo" in hint:
        return pick([
            f"Thực hiện chế độ kiểm tra, tổng hợp và báo cáo đúng hạn; ưu tiên giải pháp phù hợp với điều kiện thực tế của {tong_lop} lớp học trong nhà trường.",
            "Ban giám hiệu tổ chức kiểm tra định kỳ theo tháng, kịp thời điều chỉnh biện pháp đối với các nội dung triển khai chưa đạt yêu cầu.",
            "Các bộ phận gửi báo cáo ngắn gọn, đúng hạn; nêu rõ kết quả, khó khăn và đề xuất để nhà trường tổng hợp báo cáo cấp trên."
        ])

    fallback = [
        (
            f"Nội dung này được cụ thể hóa cho {school_name} theo hướng khả thi, xác định rõ đầu việc, "
            "đơn vị phụ trách và mốc thời gian hoàn thành."
        ),
        (
            "Các bộ phận liên quan chủ động triển khai theo chức năng, tăng cường phối hợp nội bộ và "
            "bảo đảm thống nhất trong quá trình thực hiện."
        ),
        (
            f"Quá trình triển khai phải bám sát thực tiễn học sinh, lớp học và điều kiện cơ sở vật chất của {school_name}, "
            "bảo đảm hiệu quả và tính bền vững."
        ),
    ]
    return fallback[idx % len(fallback)]


def _normalize_heading_for_thpt(line: str) -> str:
    """Chuẩn hóa tiêu đề mục để phù hợp cơ cấu trường THPT, tránh sai thẩm quyền."""
    text = line.strip()
    low = text.lower()

    if re.match(r"^1\.\s+các phòng chuyên môn thuộc", low):
        return "1. Ban giám hiệu, Văn phòng, Đoàn trường và các tổ chuyên môn"
    if re.match(r"^1\.1\.\s+văn phòng", low):
        return "1.1. Văn phòng nhà trường"
    if re.match(r"^1\.2\.\s+", low):
        return "1.2. Tổ chuyên môn, giáo viên chủ nhiệm và Đoàn trường"
    if re.match(r"^2\.\s+các đơn vị trực thuộc", low):
        return "2. Các tổ chuyên môn và bộ phận liên quan trong nhà trường"
    if re.match(r"^3\.\s+đề nghị\s+(uỷ|ủy) ban nhân dân", low):
        return "3. Phối hợp với cha mẹ học sinh và Công an địa phương"

    return text


def _rewrite_line_conservative(line: str, school_name: str) -> str:
    """Viết lại theo hướng bám sát câu gốc, chỉ thay phạm vi triển khai cho THPT."""
    text = line.strip()

    replacements = [
        ("Sở Giáo dục và Đào tạo (GDĐT)", school_name),
        ("Sở Giáo dục và Đào tạo", school_name),
        ("Sở GDĐT", school_name),
        ("Lãnh đạo Sở", "Ban giám hiệu"),
        ("lãnh đạo Sở", "ban giám hiệu"),
        ("Lãnh đạo Trường Trung học Phổ thông Đốc Binh Kiều", "Ban giám hiệu"),
        ("lãnh đạo Trường Trung học Phổ thông Đốc Binh Kiều", "ban giám hiệu"),
        ("toàn ngành", "toàn trường"),
        ("ngành Giáo dục", "nhà trường"),
        ("đơn vị trực thuộc", "tổ chuyên môn và bộ phận liên quan"),
        ("các đơn vị", "các tổ chuyên môn và bộ phận liên quan"),
        ("cơ sở giáo dục", "nhà trường"),
        ("cán bộ, công chức, viên chức, người lao động", "cán bộ, giáo viên, nhân viên"),
        ("Uỷ ban nhân dân Tỉnh", "cơ quan chức năng địa phương"),
        ("Ủy ban nhân dân Tỉnh", "cơ quan chức năng địa phương"),
    ]
    for src, dst in replacements:
        text = text.replace(src, dst)

    text = re.sub(r"tổ\s*trưởng\s*chuyên\s*môn", "tổ chuyên môn", text, flags=re.IGNORECASE)

    low = text.lower()

    if "uỷ ban nhân dân cấp xã" in low or "ủy ban nhân dân cấp xã" in low:
        return "Phối hợp với Ban đại diện cha mẹ học sinh và Công an địa phương trong tuyên truyền, nhắc nhở và xử lý các trường hợp học sinh vi phạm trật tự, an toàn giao thông."

    if "ubnd các xã" in low or "ubnd xã" in low or "xã, phường" in low:
        return "Nhà trường chủ động phối hợp chính quyền, lực lượng chức năng tại địa phương và cha mẹ học sinh để bảo đảm trật tự, an toàn giao thông cho học sinh."

    if "phòng giáo dục mầm non" in low or "giáo dục nghề nghiệp" in low:
        return "Tổ chuyên môn, giáo viên chủ nhiệm và Đoàn trường phối hợp triển khai nội dung giáo dục an toàn giao thông phù hợp học sinh trung học phổ thông."

    if "văn phòng sở" in low:
        return "Văn phòng nhà trường làm đầu mối tham mưu Ban giám hiệu chỉ đạo, theo dõi và tổng hợp kết quả thực hiện công tác bảo đảm trật tự, an toàn giao thông."

    if "thủ trưởng các cơ sở giáo dục" in low:
        return "1. Ban giám hiệu căn cứ kế hoạch này xây dựng kế hoạch chi tiết của nhà trường, phân công rõ trách nhiệm theo từng bộ phận và thời gian thực hiện."

    if "thủ trưởng các nhà trường" in low:
        return "1. Ban giám hiệu căn cứ kế hoạch này xây dựng kế hoạch chi tiết của nhà trường, phân công rõ trách nhiệm theo từng bộ phận và thời gian thực hiện."

    if "trực thuộc trên địa bàn" in low:
        return "Phối hợp với cha mẹ học sinh và lực lượng công an địa phương để triển khai hiệu quả các nội dung về bảo đảm trật tự, an toàn giao thông cho học sinh."

    if "đơn vị mình quản lý" in low:
        return text.replace("đơn vị mình quản lý", "nhà trường")

    if "các đơn vị" in low and "trực thuộc" in low:
        return text.replace("các đơn vị", "các tổ chuyên môn, giáo viên chủ nhiệm và bộ phận liên quan")

    return text


def _soan_van_ban_offline(noi_dung_van_ban_den: str, thong_tin_truong: str) -> dict:
    """Soạn nội dung cục bộ để test pipeline khi chưa dùng DeepSeek API."""
    school_ctx = _parse_context_kv(thong_tin_truong)
    school_name = school_ctx.get("ten_truong", "Trường THPT Đốc Binh Kiều")
    tong_lop = school_ctx.get("tong_lop", "13")
    tong_hoc_sinh = school_ctx.get("tong_hoc_sinh", "506")
    tong_giao_vien = school_ctx.get("tong_giao_vien", "30")

    loai_vb = _xac_dinh_loai_van_ban_tu_nguon(noi_dung_van_ban_den)
    trich_yeu = _trich_trich_yeu_tu_nguon(noi_dung_van_ban_den, loai_vb)
    can_cu_chinh = _tao_can_cu_chinh_tu_van_ban(noi_dung_van_ban_den, loai_vb, trich_yeu)

    lines = [ln.strip() for ln in (noi_dung_van_ban_den or "").splitlines() if ln.strip()]
    cleaned = []
    for i, line in enumerate(lines):
        up = line.upper()
        if i <= 2 and up in {"KẾ HOẠCH", "BÁO CÁO", "TỜ TRÌNH", "QUYẾT ĐỊNH", "BIÊN BẢN", "CÔNG VĂN"}:
            continue
        if line.lower().startswith("nơi nhận"):
            break
        if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in up:
            break
        if "KT. GIÁM ĐỐC" in up or "PHÓ GIÁM ĐỐC" in up:
            break
        if "UBND TỈNH" in up and "SỞ GIÁO DỤC" in up:
            break
        cleaned.append(line)

    rewritten = [can_cu_chinh]
    current_h1 = ""
    current_h2 = ""
    intro_opening_added = False

    for line in cleaned:
        adapted_heading = _normalize_heading_for_thpt(_adapt_line_to_thpt(line, school_name))
        if _is_heading_line(line):
            rewritten.append(adapted_heading)
            if re.match(r"^[IVXLCDM]+\.\s+", line):
                current_h1 = line
                current_h2 = ""
            elif re.match(r"^\d+(?:\.\d+)*\.\s+", line):
                current_h2 = line
            continue

        # Dòng mở đầu trước các mục chính: viết lại ngắn gọn theo bối cảnh trường.
        if not current_h1:
            line_up = line.upper()
            if line_up.startswith("CĂN CỨ"):
                # Bỏ qua vì đã có 1 dòng căn cứ chính ngắn gọn.
                continue
            elif not intro_opening_added:
                rewritten.append(_rewrite_intro_line(line, school_name))
                intro_opening_added = True
            continue

        rewritten_line = _rewrite_line_conservative(line, school_name)
        rewritten.append(rewritten_line)

    tom_tat = (
        f"Cụ thể hóa văn bản cấp trên theo bối cảnh {school_name}, giữ nguyên khung mục chính "
        "và điều chỉnh nhiệm vụ, phân công phù hợp điều kiện thực tế của nhà trường."
    )

    rewritten = _renumber_section_items(rewritten, "IV.")

    return _normalize_ai_output({
        "loai_van_ban": loai_vb,
        "so_ky_hieu_goi_y": SO_KY_HIEU_BY_TYPE.get(loai_vb, "/THPTĐBK"),
        "trich_yeu": trich_yeu,
        "tom_tat_yeu_cau": tom_tat,
        "noi_nhan": "Sở GDĐT Đồng Tháp (báo cáo); Các tổ chuyên môn; Giáo viên chủ nhiệm; Lưu: VT",
        "noi_dung": "\n".join(rewritten),
    })


def _tao_prompt_system(thong_tin_truong: str) -> str:
    """Tạo prompt hệ thống động để ép AI bám cấu trúc nguồn và bối cảnh trường."""
    return f"""
Bạn là Content Engine soạn văn bản hành chính cho THPT.
Context trường (biến dữ liệu):
{thong_tin_truong}

Ràng buộc bắt buộc:
- Chỉ trả JSON thuần với đúng 6 khóa: loai_van_ban, so_ky_hieu_goi_y, trich_yeu, tom_tat_yeu_cau, noi_nhan, noi_dung.
- Không thêm khóa, không markdown, không giải thích.
- Bảo toàn bản đồ cấu trúc văn bản Sở: thứ tự mục/phần, hệ đánh số, số lượng mục tối đa có thể.
- Phần căn cứ phải ngắn gọn: chỉ căn cứ theo văn bản chính đang xử lý (ví dụ Kế hoạch số ...), không liệt kê dài nhiều văn bản cấp trên.
- Không chèn footer vào noi_dung (Nơi nhận/ngày tháng/chữ ký/chức vụ ký); không lặp tiêu đề loại văn bản ở đầu noi_dung.
- Nội dung mỗi mục phải cụ thể theo điều kiện thật của trường; tránh câu sáo rỗng/chung chung; không bịa số liệu.
- noi_nhan là chuỗi phân tách bằng dấu chấm phẩy.

Quy tắc số ký hiệu:
- cong_van=/THPTĐBK
- ke_hoach=/KH-THPTĐBK
- bao_cao=/BC-THPTĐBK
- to_trinh=/TTr-THPTĐBK
- quyet_dinh=/QĐ-THPTĐBK
- bien_ban=/BB-THPTĐBK
"""


def goi_ai_soan_van_ban(noi_dung_van_ban_den: str, thong_tin_truong: str) -> dict:
    """Gọi DeepSeek API để soạn văn bản đi, trả về dict."""

    if not USE_DEEPSEEK_API:
        print("ℹ️  Chế độ local: không gọi DeepSeek API, soạn nội dung offline.")
        return _soan_van_ban_offline(noi_dung_van_ban_den, thong_tin_truong)

    if not DEEPSEEK_API_KEY:
        print("⚠️  Thiếu DEEPSEEK_API_KEY. Tự động fallback sang soạn offline.")
        return _soan_van_ban_offline(noi_dung_van_ban_den, thong_tin_truong)
    
    client = OpenAI(
        api_key=DEEPSEEK_API_KEY,
        base_url=DEEPSEEK_BASE_URL
    )

    prompt_system = _tao_prompt_system(thong_tin_truong)
    
    ban_do_cau_truc = _rut_ban_do_cau_truc_van_ban(noi_dung_van_ban_den)
    noi_dung_nguon_rut_gon = _rut_gon_noi_dung_nguon(noi_dung_van_ban_den)

    prompt_user = f"""
Nhiệm vụ: cụ thể hóa văn bản Sở cho bối cảnh THPT, vẫn giữ tối đa cấu trúc gốc.

Bản đồ cấu trúc nguồn:
{ban_do_cau_truc}

Nội dung nguồn đã rút gọn:
{noi_dung_nguon_rut_gon}

Trả JSON duy nhất theo schema:
{{
  "loai_van_ban": "cong_van|ke_hoach|bao_cao|to_trinh|quyet_dinh|bien_ban",
  "so_ky_hieu_goi_y": "",
  "trich_yeu": "",
  "tom_tat_yeu_cau": "",
  "noi_nhan": "",
  "noi_dung": ""
}}
"""
    
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": prompt_system},
                {"role": "user", "content": prompt_user}
            ],
            temperature=0.3,
            max_tokens=5000
        )
        
        raw = response.choices[0].message.content.strip()
        print(f"🔍 AI trả về (đầu 300 ký tự): {raw[:300]}")
        
        return _parse_json_response(raw)
    
    except Exception as e:
        print(f"❌ Lỗi gọi DeepSeek API: {e}")
        raise


def _strip_noi_dung_footer(noi_dung: str) -> str:
    """Loại bỏ footer khỏi nội dung để tránh trùng khối Nơi nhận/Ký tên trong template."""
    if not noi_dung:
        return ""

    lines = str(noi_dung).splitlines()
    cleaned = []
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            cleaned.append("")
            continue

        lower = line_stripped.lower()
        if lower.startswith("nơi nhận"):
            break
        # Chỉ cắt khi dòng ngày tháng mang dạng footer chữ ký ("Địa danh, ngày ...").
        if re.match(r"^[^,]{1,40},\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+", lower):
            break
        if line_stripped in {"./.", "TM.", "KT.", "PHÓ HIỆU TRƯỞNG", "HIỆU TRƯỞNG"}:
            break

        cleaned.append(line.rstrip())

    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    return "\n".join(cleaned)


def _remove_leading_doc_title(noi_dung: str) -> str:
    """Bỏ dòng tiêu đề loại văn bản ở đầu noi_dung nếu AI chèn vào."""
    if not noi_dung:
        return ""

    lines = str(noi_dung).splitlines()
    if not lines:
        return ""

    first = lines[0].strip().upper()
    doc_titles = {
        "KẾ HOẠCH",
        "BÁO CÁO",
        "TỜ TRÌNH",
        "QUYẾT ĐỊNH",
        "BIÊN BẢN",
        "CÔNG VĂN",
    }
    if first in doc_titles:
        lines = lines[1:]

    while lines and not lines[0].strip():
        lines = lines[1:]

    return "\n".join(lines)


def _normalize_ai_output(result: dict) -> dict:
    """Chuẩn hóa output AI: đúng 6 khóa, đúng type/số ký hiệu, nội dung sạch footer."""
    if not isinstance(result, dict):
        result = {}

    loai_van_ban = str(result.get("loai_van_ban", "cong_van")).strip().lower()
    if loai_van_ban not in ALLOWED_DOC_TYPES:
        loai_van_ban = "cong_van"

    so_ky_hieu = str(result.get("so_ky_hieu_goi_y", "")).strip()
    so_mac_dinh = SO_KY_HIEU_BY_TYPE.get(loai_van_ban, "/THPTĐBK")
    if not so_ky_hieu or not so_ky_hieu.endswith("THPTĐBK"):
        so_ky_hieu = so_mac_dinh

    trich_yeu = str(result.get("trich_yeu", "")).strip()
    tom_tat_yeu_cau = str(result.get("tom_tat_yeu_cau", "")).strip()
    noi_nhan = str(result.get("noi_nhan", "")).strip()
    noi_dung = str(result.get("noi_dung", "")).strip()

    noi_dung = _strip_noi_dung_footer(noi_dung)
    noi_dung = _remove_leading_doc_title(noi_dung)

    normalized = {
        "loai_van_ban": loai_van_ban,
        "so_ky_hieu_goi_y": so_ky_hieu,
        "trich_yeu": trich_yeu,
        "tom_tat_yeu_cau": tom_tat_yeu_cau,
        "noi_nhan": noi_nhan,
        "noi_dung": noi_dung,
    }

    # Đảm bảo tuyệt đối chỉ 6 khóa theo contract pipeline.
    return {k: normalized.get(k, "") for k in REQUIRED_JSON_KEYS}


def _parse_json_response(text: str) -> dict:
    """Parse JSON response từ AI thành dict."""
    # Tìm JSON trong response (có thể có text trước/sau)
    json_match = re.search(r'\{.*\}', text, re.DOTALL)
    
    if json_match:
        try:
            json_str = json_match.group(0)
            result = json.loads(json_str)
            return _normalize_ai_output(result)
        except json.JSONDecodeError as e:
            print(f"⚠️  Lỗi parse JSON: {e}")
            print(f"   JSON nhận được: {json_str[:200]}")
            # Fallback: trả về dict rỗng
            return _normalize_ai_output({
                'loai_van_ban': 'cong_van',
                'so_ky_hieu_goi_y': '/THPTĐBK',
                'trich_yeu': 'Văn bản không xác định',
                'tom_tat_yeu_cau': '',
                'noi_nhan': 'Sở GDĐT Đồng Tháp',
                'noi_dung': 'Lỗi parse response từ AI'
            })
    else:
        print(f"⚠️  Không tìm thấy JSON trong response")
        print(f"   Response: {text[:200]}")
        return _normalize_ai_output({
            'loai_van_ban': 'cong_van',
            'so_ky_hieu_goi_y': '/THPTĐBK',
            'trich_yeu': 'Văn bản không xác định',
            'tom_tat_yeu_cau': '',
            'noi_nhan': 'Sở GDĐT Đồng Tháp',
            'noi_dung': 'Không nhận được response từ AI'
        })


# ============================================================
# BƯỚC 3: TẠO FILE WORD THEO ĐÚNG THỂ THỨC
# ============================================================

def tao_van_ban_docx(du_lieu: dict, ten_file_goc: str) -> str:
    """
    Tạo file .docx văn bản đi theo thể thức hành chính.
    
    Orchestration layer: Chuẩn bị metadata và gọi renderer_engine.
    Không làm styling - đó là trách nhiệm của renderer_engine.
    """
    from render_ke_hoach import render_document, parse_content_to_blocks, clean_content
    
    # Tạo thư mục van-ban-di nếu chưa có
    duong_dan_thu_muc = Path(__file__).parent / THU_MUC_VAN_BAN_DI
    duong_dan_thu_muc.mkdir(parents=True, exist_ok=True)
    
    # Chuẩn bị metadata từ AI output
    loai = du_lieu.get("loai_van_ban", "cong_van")
    loai_key = str(loai).strip().lower()
    loai_hien_thi = {
        "cong_van": "CÔNG VĂN",
        "ke_hoach": "KẾ HOẠCH",
        "bao_cao": "BÁO CÁO",
        "to_trinh": "TỜ TRÌNH",
        "quyet_dinh": "QUYẾT ĐỊNH",
        "bien_ban": "BIÊN BẢN",
    }.get(loai_key, str(loai).strip().upper() if loai else "KẾ HOẠCH")
    so_ky_hieu = du_lieu.get("so_ky_hieu_goi_y", SO_KY_HIEU_BY_TYPE.get(loai_key, "/THPTĐBK"))
    trich_yeu = du_lieu.get("trich_yeu", "")
    noi_dung  = du_lieu.get("noi_dung", "")
    noi_nhan_raw = du_lieu.get("noi_nhan", "")
    noi_nhan = dinh_dang_noi_nhan(noi_nhan_raw, noi_dung)
    
    # Parse nội dung thành structured blocks (AI Parser layer)
    content_clean = clean_content(noi_dung)
    blocks = parse_content_to_blocks(content_clean)
    
    # Chuẩn bị metadata cho renderer
    metadata = {
        'loai_van_ban': loai_hien_thi,
        'so_ky_hieu': so_ky_hieu,
        'ngay_thang': du_lieu.get('ngay_thang', dinh_dang_ngay_thang_hien_tai()),
        'trich_yeu': trich_yeu,
        'noi_nhan': noi_nhan,
        'nguoi_ky': 'Nguyễn Minh Trí',  # Có thể lấy từ du_lieu
        'chuc_vu_ky': 'KT. HIỆU TRƯỞNG\nPHÓ HIỆU TRƯỞNG'  # Có thể lấy từ du_lieu
    }
    
    # Lấy template path
    template_path = Path(__file__).parent / "TEMPLATE.docx"
    
    # Đặt tên file - thêm timestamp để tránh trùng
    from datetime import datetime
    timestamp = datetime.now().strftime("%H%M%S")
    ten_file_cu = Path(ten_file_goc).stem
    ten_file_moi = f"VBDi_{ten_file_cu}_{loai}_{timestamp}.docx"
    duong_dan_luu = duong_dan_thu_muc / ten_file_moi
    
    # Render document (Renderer Engine layer)
    render_document(
        template_path=template_path,
        output_path=duong_dan_luu,
        metadata=metadata,
        blocks=blocks
    )
    
    return str(duong_dan_luu)


# ============================================================
# HÀM SOẠN BÁO CÁO (nếu có yêu cầu)
# ============================================================

def tao_bao_cao_docx(du_lieu: dict, ten_file_goc: str) -> str:
    """Tạo file báo cáo riêng nếu văn bản đến có yêu cầu báo cáo."""
    noi_dung_bc = du_lieu.get("noi_dung_bao_cao", "")
    if not noi_dung_bc:
        return None
    
    # Dùng lại hàm tạo docx nhưng với dữ liệu báo cáo
    du_lieu_bc = {
        "loai_van_ban": "bao_cao",
        "so_ky_hieu_goi_y": "/BC-THPTĐBK",
        "trich_yeu": du_lieu.get("trich_yeu", ""),
        "noi_dung": noi_dung_bc,
        "noi_nhan": du_lieu.get("noi_nhan", ""),
        "can_bao_cao": False,
    }
    ten_file_bc = Path(ten_file_goc).stem + "_baocao"
    return tao_van_ban_docx(du_lieu_bc, ten_file_bc)


# ============================================================
# HÀM CHÍNH: XỬ LÝ 1 FILE VĂN BẢN ĐẾN
# ============================================================

def xu_ly_mot_van_ban(duong_dan_file: str) -> dict:
    """
    Xử lý một file văn bản đến:
    1. Đọc nội dung
    2. Gọi AI soạn
    3. Tạo file Word
    4. Trả về kết quả
    """
    ten_file = Path(duong_dan_file).name
    print(f"\n{'='*60}")
    print(f"📄 Đang xử lý: {ten_file}")
    print(f"{'='*60}")
    
    # Bước 1: Đọc nội dung
    print("⏳ Đọc nội dung file...")
    noi_dung = doc_noi_dung_file(duong_dan_file)
    if not noi_dung or len(noi_dung) < 100:
        print(f"⚠️  File rỗng hoặc không đọc được: {ten_file}")
        return {"loi": "Không đọc được file"}
    
    # Bước 2: Gọi AI
    print("🤖 Đang nhờ DeepSeek AI soạn văn bản...")
    try:
        context_truong = _lay_context_truong_toi_uu()
        du_lieu = goi_ai_soan_van_ban(noi_dung, context_truong)
    except json.JSONDecodeError as e:
        print(f"❌ Lỗi parse JSON từ AI: {e}")
        return {"loi": f"Lỗi parse JSON: {e}"}
    except Exception as e:
        print(f"❌ Lỗi gọi API: {e}")
        return {"loi": str(e)}
    
    # In tóm tắt
    print(f"\n📋 Tóm tắt yêu cầu:")
    print(f"   {du_lieu.get('tom_tat_yeu_cau', '')}")
    print(f"\n📝 Loại văn bản đi: {du_lieu.get('loai_van_ban', '').upper()}")
    print(f"   Trích yếu: {du_lieu.get('trich_yeu', '')}")
    
    # Bước 3: Tạo file Word
    print("\n💾 Đang tạo file Word...")
    ket_qua = {"ten_file_goc": ten_file}
    
    try:
        duong_dan_vbd = tao_van_ban_docx(du_lieu, duong_dan_file)
        ket_qua["van_ban_di"] = duong_dan_vbd
        print(f"✅ Đã tạo: {duong_dan_vbd}")
    except Exception as e:
        print(f"❌ Lỗi tạo file Word: {e}")
        ket_qua["loi_van_ban_di"] = str(e)
    
    # Bước 4: Không tạo báo cáo (để tiết kiệm token)
    
    return ket_qua


# ============================================================
# XỬ LÝ NHIỀU FILE (TẤT CẢ TRONG THƯ MỤC)
# ============================================================

def xu_ly_thu_muc(thu_muc: str = None) -> list:
    """Xử lý tất cả file trong thư mục văn bản đến."""
    thu_muc = thu_muc or THU_MUC_VAN_BAN_DEN
    
    # Chuyển đường dẫn tương đối thành tuyệt đối
    duong_dan_thu_muc = Path(__file__).parent / thu_muc
    
    if not duong_dan_thu_muc.exists():
        print(f"❌ Thư mục không tồn tại: {duong_dan_thu_muc}")
        return []
    
    files = [
        str(duong_dan_thu_muc / f)
        for f in os.listdir(duong_dan_thu_muc)
        if f.lower().endswith((".docx", ".pdf"))
    ]
    
    if not files:
        print(f"📂 Không có file nào trong: {thu_muc}")
        return []
    
    print(f"\n🗂️  Tìm thấy {len(files)} file cần xử lý trong '{duong_dan_thu_muc}'")
    
    ket_qua_tat_ca = []
    for f in sorted(files):
        kq = xu_ly_mot_van_ban(f)
        ket_qua_tat_ca.append(kq)
    
    # Tổng kết
    thanh_cong = sum(1 for k in ket_qua_tat_ca if "van_ban_di" in k)
    co_bao_cao = sum(1 for k in ket_qua_tat_ca if "bao_cao" in k)
    
    print(f"\n{'='*60}")
    print(f"🎉 HOÀN THÀNH!")
    print(f"   ✅ Soạn thành công: {thanh_cong}/{len(files)} văn bản")
    print(f"   📊 Có báo cáo kèm: {co_bao_cao} văn bản")
    print(f"   📁 File lưu tại: {duong_dan_thu_muc.parent / 'van-ban-di'}/")
    print(f"{'='*60}")
    
    return ket_qua_tat_ca


# ============================================================
# CHẠY TRỰC TIẾP (test với 1 file)
# ============================================================
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        # Xử lý file cụ thể
        duong_dan = sys.argv[1]
        xu_ly_mot_van_ban(duong_dan)
    else:
        # Xử lý toàn bộ thư mục
        xu_ly_thu_muc()
