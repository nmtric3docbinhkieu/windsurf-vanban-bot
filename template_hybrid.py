#!/usr/bin/env python3
"""
Script kết hợp: docxtpl cho header/footer + python-docx cho nội dung
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate
from pathlib import Path
import re

def clean_content(content):
    """Lọc bỏ tiêu đề và footer"""
    lines = content.split('\n')
    
    # Bỏ dòng đầu tiên nếu là loại văn bản
    if lines and lines[0].strip().upper() in ["KẾ HOẠCH", "QUYẾT ĐỊNH", "BÁO CÁO", "TỜ TRÌNH", "BIÊN BẢN", "CÔNG VĂN"]:
        lines = lines[1:]
    
    # Bỏ dòng thứ hai nếu là trích yếu
    if lines and lines[0].strip():
        lines = lines[1:]
    
    # Bỏ dòng trống tiếp theo
    if lines and not lines[0].strip():
        lines = lines[1:]
    
    # Lọc nội dung chính, bỏ footer
    result = []
    for line in lines:
        line = line.strip()
        
        # Tìm dấu hiệu footer
        if "Nơi nhận:" in line or re.match(r"^[^,]+,\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+", line, re.IGNORECASE):
            break
        
        # Loại bỏ "./." ở cuối
        line = line.replace("./.", "").strip()
        
        if line:
            result.append(line)
        else:
            result.append('')
    
    return '\n'.join(result)

def add_paragraph_formatted(doc, text, bold=False, indent=0):
    """Thêm đoạn văn với định dạng đầy đủ"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    
    if indent > 0:
        p.paragraph_format.first_line_indent = Cm(indent)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = bold
    
    # Set font tiếng Việt
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.append(rFonts)

def add_content_formatted(doc, content):
    """Thêm nội dung với định dạng đẹp"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Đề mục cấp I, II, III - in đậm, không thụt đầu dòng
        if re.match(r'^[IVX]+\.', line):
            add_paragraph_formatted(doc, line, bold=True, indent=0)
        # Đề mục cấp 1, 2, 3 - in đậm, thụt đầu dòng 1cm
        elif re.match(r'^\d+\.', line):
            add_paragraph_formatted(doc, line, bold=True, indent=1.0)
        # Đề mục cấp a, b, c - bình thường, thụt đầu dòng 1.5cm
        elif re.match(r'^[a-z]+\.', line):
            add_paragraph_formatted(doc, line, bold=False, indent=1.5)
        # Bullet
        elif line.startswith('- '):
            clean = line[2:]
            add_paragraph_formatted(doc, '• ' + clean, bold=False, indent=0.75)
        # Dòng thường - thụt đầu dòng 1.25cm
        else:
            add_paragraph_formatted(doc, line, bold=False, indent=1.25)

def create_van_ban_hybrid(template_path, content, output_path, metadata):
    """
    Tạo văn bản kết hợp: docxtpl cho placeholder + python-docx cho nội dung
    
    Args:
        template_path: File template có header/footer chuẩn
        content: Nội dung chính
        output_path: File output
        metadata: Dict với so_ky_hieu, ngay_thang, trich_yeu, noi_nhan, nguoi_ky, chuc_vu_ky
    """
    # Bước 1: Inject placeholder bằng docxtpl
    doc_tpl = DocxTemplate(template_path)
    
    # Chuẩn bị data cho placeholder (không có noi_dung)
    data = {
        'so_ky_hieu': metadata.get('so_ky_hieu', ''),
        'ngay_thang': metadata.get('ngay_thang', ''),
        'trich_yeu': metadata.get('trich_yeu', ''),
        'noi_nhan': metadata.get('noi_nhan', ''),
        'nguoi_ky': metadata.get('nguoi_ky', ''),
        'chuc_vu_ky': metadata.get('chuc_vu_ky', '')
    }
    
    # Render template
    doc_tpl.render(data)
    
    # Bước 2: Lưu temp file
    temp_path = output_path.parent / f"temp_{output_path.name}"
    doc_tpl.save(temp_path)
    
    # Bước 3: Mở lại bằng python-docx để thêm nội dung
    doc = Document(temp_path)
    
    # Tìm vị trí placeholder {{noi_dung}} và thay bằng nội dung thực
    content = clean_content(content)
    
    # Tìm paragraph có placeholder
    placeholder_index = -1
    for i, para in enumerate(doc.paragraphs):
        if '{{noi_dung}}' in para.text:
            placeholder_index = i
            break
    
    if placeholder_index >= 0:
        # Xóa placeholder
        doc.paragraphs[placeholder_index].clear()
        # Thêm nội dung formatted vào đúng vị trí
        # Cách: xóa tất cả paragraph sau placeholder, sau đó thêm nội dung, rồi thêm lại các paragraph bị xóa
        # Cách đơn giản hơn: tạo paragraph mới tại vị trí placeholder
        
        # Lấy tất cả paragraph sau placeholder
        after_paragraphs = []
        for i in range(placeholder_index + 1, len(doc.paragraphs)):
            para = doc.paragraphs[placeholder_index + 1]
            after_paragraphs.append(para)
            # Không xóa, sẽ xử lý sau
        
        # Xóa các paragraph sau placeholder
        for _ in range(len(doc.paragraphs) - placeholder_index - 1):
            doc.paragraphs[placeholder_index + 1]._element.getparent().remove(doc.paragraphs[placeholder_index + 1]._element)
        
        # Thêm nội dung formatted
        add_content_formatted(doc, content)
        
        # Thêm lại các paragraph bị xóa (nơi nhận, chữ ký)
        for para in after_paragraphs:
            if para.text.strip():  # Chỉ thêm paragraph có nội dung
                doc.add_paragraph(para.text)
    else:
        # Nếu không tìm thấy placeholder, thêm vào cuối
        add_content_formatted(doc, content)
    
    # Save
    doc.save(output_path)
    
    # Xóa temp file
    temp_path.unlink()
    
    print(f"✅ Đã tạo file: {output_path}")

def test():
    """Test với dữ liệu mẫu"""
    template_path = Path(__file__).parent / "TEMPLATE.docx"
    content_file = Path(__file__).parent / "noi_dung_ke_hoach.txt"
    output_path = Path(__file__).parent.parent / "van-ban-di" / "test_hybrid_v2.docx"
    
    content = content_file.read_text(encoding='utf-8')
    
    metadata = {
        'so_ky_hieu': '123/KH-THPTĐBK',
        'ngay_thang': 'Đốc Bình Kiều, ngày 10 tháng 5 năm 2026',
        'trich_yeu': 'Triển khai thực hiện Chỉ thị số 03/CT-TTg về ngăn chặn bạo lực học đường',
        'noi_nhan': 'Sở GDĐT Đồng Tháp (báo cáo); Lưu: VT',
        'nguoi_ky': 'Nguyễn Minh Trí',
        'chuc_vu_ky': 'KT. HIỆU TRƯỞNG\nPHÓ HIỆU TRƯỞNG'
    }
    
    create_van_ban_hybrid(template_path, content, output_path, metadata)

if __name__ == "__main__":
    test()
